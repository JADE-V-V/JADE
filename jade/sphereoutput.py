# -*- coding: utf-8 -*-
"""
Created on Thu Jan  2 10:36:38 2020

@author: Davide Laghi

Copyright 2021, the JADE Development Team. All rights reserved.

This file is part of JADE.

JADE is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

JADE is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with JADE.  If not, see <http://www.gnu.org/licenses/>.
"""

from __future__ import annotations

import abc
import itertools
import json
import logging
import math
import os
import shutil
from typing import TYPE_CHECKING

import numpy as np
import pandas as pd
from docx.shared import Inches
from tqdm import tqdm
from xlsxwriter.utility import xl_rowcol_to_cell

import jade.atlas as at
import jade.excelsupport as exsupp
import jade.plotter as plotter
from jade.output import AbstractBenchmarkOutput, MCNPSimOutput, OpenMCSimOutput
from f4enix.output.mctal import Mctal

if TYPE_CHECKING:
    from jade.main import Session

from jade.__openmc__ import OMC_AVAIL

if OMC_AVAIL:
    import jade.openmc as omc


class AbstractSphereBenchmarkOutput(AbstractBenchmarkOutput):
    def __init__(self, lib: str, code: str, testname: str, session: Session) -> None:
        """
        Initialises the SphereOutput class from the general BenchmarkOutput
        class, see output.py for details on how self variables are assigned

        Parameters
        ----------
        lib : str
            library to post-process
        code : str
            code being post processed
        testname : str
            name of the benchmark being postprocessed
        session : Session
            Jade Session
        exp : str
            the benchmark is an experimental one

        Returns
        -------
        None.

        """
        super().__init__(lib, code, testname, session)

        # Load the settings for zaids and materials
        mat_path = os.path.join(self.cnf_path, "MaterialsSettings.csv")
        self.mat_settings = pd.read_csv(mat_path, sep=",").set_index("Symbol")

        zaid_path = os.path.join(self.cnf_path, "ZaidSettings.csv")
        self.zaid_settings = pd.read_csv(zaid_path, sep=",").set_index("Z")

        # # The metadata needs to be re-read since no multitest is foreseen in the
        # # normal BenchmarkOutput class
        # # Read the metadata, they should be all equal
        # try:
        #     results_path = os.path.join(
        #         self.test_path, os.listdir(self.test_path)[0], code
        #     )
        #     self.metadata = self._read_metadata_run(results_path)
        # except TypeError:
        #     # means that self.test_path is a dict, hence a comparison. No
        #     # metadata involved here
        #     self.metadata = None

    def _get_output_files(self, results_path: str | os.PathLike, code: str) -> list:
        """
        Enforced method from inheritance of AbstractBenchmarkOutput, not used in Sphere.

        Parameters
        ----------
        results_path : str | os.PathLike
            Path to simulation results folder.

        Returns
        -------
        list
            List of simulation results files.
        """
        pass

    def parse_output_data(self, results_path: str | os.PathLike):
        """
        Abstract function for retrieving simulation output data, tally numbers and tally comments.
        Not used in Sphere.

        Parameters
        ----------
        results_path : str | os.PathLike
            Path to simulation results

        Returns
        -------
        None
        """

    def single_postprocess(self) -> None:
        """
        Execute the full post-processing of a single library (i.e. excel,
        raw data and atlas)

        Returns
        -------
        None.

        """
        print(" Generating Excel Recap...")
        self.pp_excel_single()
        print(" Dumping Raw Data...")
        self.print_raw()
        print(" Generating plots...")
        self._generate_single_plots()

    @abc.abstractmethod
    def _read_output(self) -> tuple[dict, list, list, list | None]:
        """
        Reads all outputs for a library. To be implemented for each different code.

        Returns
        -------
        outputs : dic
            Dictionary of sphere output objects used in plotting, keys are material name or ZAID number
        results : dic
            Dictionary of overview of Tally values for each material/ZAID, returns either all values > 0 for
            tallies with postive values only, all Values = 0 for empty tallies, and returns the corresponding
            tally bin if it finds any negative values. Contents of the "Values" worksheet.
        errors : dic
            Dictionary of average errors for each tally for each material/Zaid. Contents of the "Errors" worksheet.
        stat_checks : dic
            Dictionary the MCNP statistical check results for each material/ZAID. Contents of the "Statistical
            Checks" Worksheet.
        """
        pass

    def _generate_single_plots(self) -> None:
        """
        Generate all the requested plots in a temporary folder

        Returns
        -------
        None.

        """

        # edited by T. Wheeler. openmc requires separate tally numbers which is accounted for here
        # TODO this should be brought into the implementation, but actually the tally numbers
        # will be made the same across codes so this will not be needed.
        outpath = os.path.join(self.atlas_path, "tmp")
        os.mkdir(outpath)
        if self.openmc:
            tally_info = [
                (
                    4,
                    "Averaged Neutron Flux (175 groups)",
                    "Neutron Flux",
                    r"$\#/cm^2$",
                ),
                (14, "Averaged Gamma Flux (24 groups)", "Gamma Flux", r"$\#/cm^2$"),
            ]
        else:
            tally_info = [
                (
                    2,
                    "Averaged Neutron Flux (175 groups)",
                    "Neutron Flux",
                    r"$\#/cm^2$",
                ),
                (32, "Averaged Gamma Flux (24 groups)", "Gamma Flux", r"$\#/cm^2$"),
            ]
        for tally, title, quantity, unit in tally_info:
            print(" Plotting tally n." + str(tally))
            for zaidnum, output in tqdm(self.outputs.items()):
                tally_data = output.tallydata.set_index("Tally N.").loc[tally]
                energy = tally_data["Energy"].values
                values = tally_data["Value"].values
                error = tally_data["Error"].values
                lib_name = self.session.conf.get_lib_name(self.lib)
                lib = {
                    "x": energy,
                    "y": values,
                    "err": error,
                    "ylabel": str(zaidnum) + " (" + lib_name + ")",
                }
                data = [lib]
                outname = str(zaidnum) + "-" + self.lib + "-" + str(tally)
                plot = plotter.Plotter(
                    data,
                    title,
                    outpath,
                    outname,
                    quantity,
                    unit,
                    "Energy [MeV]",
                    self.testname,
                )
                plot.plot("Binned graph")

        self._build_atlas(outpath)

    def _build_atlas(self, outpath: str | os.PathLike) -> None:
        """
        Build the atlas using all plots contained in directory

        Parameters
        ----------
        outpath : str or path
            temporary folder containing all plots.

        Returns
        -------
        None.

        """
        atlas_path = os.path.join(outpath, "..")
        # Printing Atlas
        template = os.path.join(self.path_templates, "AtlasTemplate.docx")
        if self.single:
            name = self.lib
        else:
            name = self.name

        atlas = at.Atlas(template, "Sphere " + name)
        atlas.build(outpath, self.session.lib_manager, self.mat_settings)
        atlas.save(atlas_path)
        # Remove tmp images
        shutil.rmtree(outpath)

    def compare(self) -> None:
        """
        Execute the full post-processing of a comparison of libraries
        (i.e. excel, and atlas)

        Returns
        -------
        None.

        """
        print(" Generating Excel Recap...")
        self.pp_excel_comparison()
        print(" Creating Atlas...")
        # outpath = os.path.join(self.atlas_path, 'tmp')
        # os.mkdir(outpath)
        # Recover all libraries and zaids involved
        libraries, allzaids, outputs = self._get_organized_output()

        globalname = ""
        for lib in self.lib:
            globalname = globalname + lib + "_Vs_"

        globalname = globalname[:-4]

        # Plot everything
        print(" Generating Plots Atlas...")
        self._generate_plots(allzaids, globalname)

    def _generate_plots(self, allzaids: list, globalname: str) -> None:
        """
        Generate all the plots requested by the Sphere leakage benchmark

        Parameters
        ----------
        allzaids : list
            list of all zaids resulting from the union of the results from
            both libraries.
        globalname : str
            name for the output.

        Returns
        -------
        None.

        """
        outpath = os.path.join(self.atlas_path, "tmp")
        if not os.path.exists(outpath):
            os.mkdir(outpath)
        # TODO this if else should be removed once the tally numbers are made the same
        if self.code == "mcnp":
            tally_info = [
                (
                    2,
                    "Averaged Neutron Flux (175 groups)",
                    "Neutron Flux",
                    r"$\#/cm^2$",
                ),
                (32, "Averaged Gamma Flux (24 groups)", "Gamma Flux", r"$\#/cm^2$"),
            ]
        if self.code == "openmc":
            tally_info = [
                (
                    4,
                    "Averaged Neutron Flux (175 groups)",
                    "Neutron Flux",
                    r"$\#/cm^2$",
                ),
                (14, "Averaged Gamma Flux (24 groups)", "Gamma Flux", r"$\#/cm^2$"),
            ]
        for tally, title, quantity, unit in tally_info:
            print(" Plotting tally n." + str(tally))
            for zaidnum in tqdm(allzaids):
                # title = title
                data = []
                for library, lib_outputs in self.outputs.items():
                    try:  # Zaid could not be common to the libraries
                        tally_data = (
                            lib_outputs[zaidnum]
                            .tallydata.set_index("Tally N.")
                            .loc[tally]
                        )
                        # print(lib_outputs[zaidnum])
                        energy = tally_data["Energy"].values
                        values = tally_data["Value"].values
                        error = tally_data["Error"].values
                        lib_name = self.session.conf.get_lib_name(library)
                        lib = {
                            "x": energy,
                            "y": values,
                            "err": error,
                            "ylabel": str(zaidnum) + " (" + str(lib_name) + ")",
                        }
                        data.append(lib)
                    except KeyError:
                        # It is ok, simply nothing to plot here
                        pass

                outname = str(zaidnum) + "-" + globalname + "-" + str(tally)
                plot = plotter.Plotter(
                    data,
                    title,
                    outpath,
                    outname,
                    quantity,
                    unit,
                    "Energy [MeV]",
                    self.testname,
                )
                try:
                    plot.plot("Binned graph")
                except IndexError:
                    print(data)

        self._build_atlas(outpath)

    def _get_organized_output(self) -> tuple[list, list, list]:
        """
        Organizes the outputs for each library in each code in the
        outputs object

        Returns:
        --------
        libraries: list
            list of all libraries to be post processed
        allzaids: list
            list of all zaids/materials that have been run
        outputs: list
            list of all output objects for all codes and all libraries

        """
        libraries = []
        outputs = []
        zaids = []

        for libname, outputslib in self.outputs.items():
            libraries.append(libname)
            outputs.append(outputslib)
            zaids.append(list(outputslib.keys()))
        # Extend list to all zaids
        allzaids = zaids[0]
        for zaidlist in zaids[1:]:
            allzaids.extend(zaidlist)
        allzaids = set(allzaids)  # no duplicates

        return libraries, allzaids, outputs

    def _generate_dataframe(
        self, results: dict, errors: dict, stat_checks: dict | None = None
    ) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
        """
        Function to turn the output of the read_{code}_output functions into DataFrames
        for use with xlsxwriter

        Parameters
        ----------
        results : dict
            dictionary of tally summaries for each material/ZAID.
        errors : dict
            dictionaty of average tally errors across all energy bins.
        stat_checks : dict or None
            dictionary containing results of MCNP statistical checks
            (MCNP only). Defaults to None.

        Returns
        -------
        results : pd.DataFrame
            previous dictionary but in DataFrame form
        errors : pd.DataFrame
            previous dictionary but in DataFrame form
        stat_checks : pd.DataFrame
            previous dictionary but in DataFrame form
        """
        # Generate DataFrames
        results = pd.DataFrame(results)
        errors = pd.DataFrame(errors)

        # Swap Columns and correct zaid sorting
        # results
        for df in [results, errors]:
            df["index"] = pd.to_numeric(df["Zaid"].values, errors="coerce")
            df.sort_values("index", inplace=True)
            del df["index"]

            df.set_index(["Zaid", "Zaid/Mat Name"], inplace=True)
            df.reset_index(inplace=True)

        if stat_checks is not None:
            stat_checks = pd.DataFrame(stat_checks)
            stat_checks["index"] = pd.to_numeric(
                stat_checks["Zaid"].values, errors="coerce"
            )
            stat_checks.sort_values("index", inplace=True)
            del stat_checks["index"]

            stat_checks.set_index(["Zaid", "Zaid/Mat Name"], inplace=True)
            stat_checks.reset_index(inplace=True)
        return results, errors, stat_checks

    def pp_excel_single(self) -> None:
        """
        Generate the single library results excel

        Returns
        -------
        None.

        """
        outfolder_path = self.excel_path
        # os.makedirs(outfolder_path, exist_ok=True)
        # outpath = os.path.join(self.excel_path_mcnp,'Sphere_single_' + 'MCNP_' + self.lib+'.xlsx')
        outpath = os.path.join(
            outfolder_path, f"Sphere_single_{self.code}_{self.lib}.xlsx"
        )
        outputs, results, errors, stat_checks = self._read_output()
        results, errors, stat_checks = self._generate_dataframe(
            results, errors, stat_checks
        )
        self.outputs = outputs
        self.results = results
        self.errors = errors
        self.stat_checks = stat_checks
        lib_name = self.session.conf.get_lib_name(self.lib)
        # Generate DataFrames
        # results = pd.DataFrame(results)
        # errors = pd.DataFrame(errors)
        # stat_checks = pd.DataFrame(stat_checks)

        # Swap Columns and correct zaid sorting
        # results
        # for df in [results, errors, stat_checks]:
        #    df['index'] = pd.to_numeric(df['Zaid'].values, errors='coerce')
        #    df.sort_values('index', inplace=True)
        #    del df['index']

        #    df.set_index(['Zaid', 'Zaid Name'], inplace=True)
        #    df.reset_index(inplace=True)
        exsupp.sphere_single_excel_writer(
            self, outpath, lib_name, results, errors, stat_checks
        )

    @abc.abstractmethod
    def _get_output(self, results_path: str) -> SphereTallyOutput:
        """
        Get the output files for the code being post-processed.

        Returns
        -------
        AbstractSimulationOutput
        """
        pass

    def pp_excel_comparison(self) -> None:
        """
        Compute the data and create the excel for all libraries comparisons.
        In the meantime, additional data is stored for future plots.


        Returns
        -------
        None.

        """

        iteration = 0
        outputs = {}
        for reflib, tarlib, name in self.couples:
            outfolder_path = self.excel_path
            # os.mkdir(outfolder_path)
            outpath = os.path.join(
                outfolder_path, f"Sphere_comparison_{name}_{self.code}.xlsx"
            )
            # outpath = os.path.join(self.excel_path_mcnp, 'Sphere_comparison_' +
            #                       name+'.xlsx')
            # Get results
            comp_dfs = []
            error_dfs = []

            for test_path in [
                self.test_path[reflib],
                self.test_path[tarlib],
            ]:
                results = []
                errors = []
                iteration = iteration + 1
                outputs_lib = {}
                for folder in os.listdir(test_path):
                    results_path = os.path.join(test_path, folder, self.code)
                    pieces = folder.split("_")
                    # Get zaid
                    zaidnum = pieces[-2]
                    # Check for material exception
                    if zaidnum == "Sphere":
                        zaidnum = pieces[-1].upper()
                        zaidname = self.mat_settings.loc[zaidnum, "Name"]
                    else:
                        zaidname = pieces[-1]

                    output = self._get_output(results_path)

                    outputs_lib[zaidnum] = output

                    # TODO to remove when tallies are the same
                    if self.code == "mcnp":
                        res, err, columns = output.get_comparison_data(
                            ["12", "22", "24", "14", "34", "6", "46"], "mcnp"
                        )
                    elif self.code == "openmc":
                        res, err, columns = output.get_comparison_data(
                            ["4", "14"], "openmc"
                        )

                    try:
                        zn = int(zaidnum)
                    except ValueError:  # Happens for typical materials
                        zn = zaidnum

                    res.append(zn)
                    err.append(zn)
                    res.append(zaidname)
                    err.append(zaidname)

                    results.append(res)
                    errors.append(err)

                # Add reference library outputs
                if iteration == 1:
                    outputs[reflib] = outputs_lib

                if iteration == 2:
                    outputs[tarlib] = outputs_lib

                # Generate DataFrames
                columns.extend(["Zaid", "Zaid/Mat Name"])
                comp_df = pd.DataFrame(results, columns=columns)
                error_df = pd.DataFrame(errors, columns=columns)
                comp_df.set_index(["Zaid", "Zaid/Mat Name"], inplace=True)
                error_df.set_index(["Zaid", "Zaid/Mat Name"], inplace=True)
                comp_dfs.append(comp_df)
                error_dfs.append(error_df)

            self.outputs = outputs
            # self.results["mcnp"] = results
            # self.errors["mcnp"] = errors

            # Consider only common zaids
            idx1 = comp_dfs[0].index
            idx2 = comp_dfs[1].index
            newidx = idx1.intersection(idx2)

            # Build the final excel data
            final = (comp_dfs[0].loc[newidx] - comp_dfs[1].loc[newidx]) / comp_dfs[
                0
            ].loc[newidx]
            absdiff = comp_dfs[0].loc[newidx] - comp_dfs[1].loc[newidx]

            # self.diff_data["mcnp"] = final
            # self.absdiff["mcnp"] = absdiff

            # Standard deviation
            idx1 = absdiff.index
            idx2 = error_dfs[0].index
            newidx = idx1.intersection(idx2)

            std_dev = absdiff.loc[newidx] / (
                error_dfs[0].loc[newidx] * comp_dfs[0].loc[newidx]
            )

            # self.std_dev["mcnp"] = std_dev
            # Correct sorting
            for df in [final, absdiff, std_dev]:
                df.reset_index(inplace=True)
                df["index"] = pd.to_numeric(df["Zaid"].values, errors="coerce")
                df.sort_values("index", inplace=True)
                del df["index"]
                df.set_index(["Zaid", "Zaid/Mat Name"], inplace=True)

            # Create and concat the summary
            old_l = 0
            old_lim = 0
            rows = []
            limits = [0, 0.05, 0.1, 0.2, 0.2]
            for i, sup_lim in enumerate(limits[1:]):
                if i == len(limits) - 2:
                    row = {"Range": "% of cells > " + str(sup_lim * 100)}
                    for column in final.columns:
                        cleaned = final[column].replace("", np.nan).dropna()
                        l_range = len(cleaned[abs(cleaned) > sup_lim])
                        try:
                            row[column] = l_range / len(cleaned)
                        except ZeroDivisionError:
                            row[column] = np.nan
                else:
                    row = {
                        "Range": str(old_lim * 100)
                        + " < "
                        + "% of cells"
                        + " < "
                        + str(sup_lim * 100)
                    }
                    for column in final.columns:
                        cleaned = final[column].replace("", np.nan).dropna()
                        lenght = len(cleaned[abs(cleaned) < sup_lim])
                        old_l = len(cleaned[abs(cleaned) < limits[i]])
                        l_range = lenght - old_l
                        try:
                            row[column] = l_range / len(cleaned)
                        except ZeroDivisionError:
                            row[column] = np.nan

                old_lim = sup_lim
                rows.append(row)

            summary = pd.DataFrame(rows)
            summary.set_index("Range", inplace=True)
            # If it is zero the CS are equal! (NaN if both zeros)
            for df in [final, absdiff, std_dev]:
                # df[df == np.nan] = 'Not Available'
                df.astype({col: float for col in df.columns[1:]})
                df.replace(np.nan, "Not Available", inplace=True)
                df.replace(float(0), "Identical", inplace=True)
                df.replace(-np.inf, "Reference = 0", inplace=True)
                df.replace(1, "Target = 0", inplace=True)

            # retrieve single pp files to add as extra tabs to comparison workbook
            single_pp_files = []
            # Add single pp sheets
            for lib in [reflib, tarlib]:
                pp_dir = self.session.state.get_path(
                    "single", [lib, "Sphere", self.code, "Excel"]
                )
                pp_file = os.listdir(pp_dir)[0]
                single_pp_path = os.path.join(pp_dir, pp_file)
                single_pp_files.append(single_pp_path)

            # --- Write excel ---
            # Generate the excel
            exsupp.sphere_comp_excel_writer(
                self,
                outpath,
                name,
                final,
                absdiff,
                std_dev,
                summary,
                single_pp_files,
            )

            # # Add single pp sheets
            # current_wb = openpyxl.load_workbook(outpath)
            # for lib in [reflib, tarlib]:
            #     cp = self.session.state.get_path(
            #         "single", [lib, "Sphere", "mcnp", "Excel"]
            #     )
            #     file = os.listdir(cp)[0]
            #     cp = os.path.join(cp, file)
            #     # open file
            #     single_wb = openpyxl.load_workbook(cp)
            #     for ws in single_wb.worksheets:
            #         destination = current_wb.create_sheet(ws.title + " " + lib)
            #         exsupp.copy_sheet(ws, destination)
            #     single_wb.close()

            # current_wb.save(outpath)
            # current_wb.close()

            # ex.save()
            # ""

    def print_raw(self) -> None:
        """
        Assigns a path and prints the post processing data as a .csv

        Returns
        -------
        None

        """
        for key, data in self.raw_data.items():
            file = os.path.join(self.raw_path, self.code + key + ".csv")
            data.to_csv(file, header=True, index=False)

        metadata_file = os.path.join(self.raw_path, "metadata.json")
        with open(metadata_file, "w", encoding="utf-8") as outfile:
            json.dump(self.metadata, outfile, indent=4)

    def _read_metadata_run(self, simulation_folder: os.PathLike) -> dict:
        """
        Retrieve the metadata from the run

        Parameters
        ----------
        pathtofile : os.PathLike
            path to metadata file

        Returns
        -------
        metadata : dict
            metadata dictionary
        """
        # the super can be used, just changing the expected path
        dirname = os.path.dirname(simulation_folder)
        folder = os.path.join(
            dirname,
            os.listdir(dirname)[0],
            self.code,
        )

        metadata = super()._read_metadata_run(folder)

        return metadata


class MCNPSphereBenchmarkOutput(AbstractSphereBenchmarkOutput):
    def _read_code_version(self, simulation_folder: str | os.PathLike) -> str | None:
        """
        Function to retrieve MCNP code version. Implimentation should be added to child classes for each code.

        Parameters
        ----------
        simulation_folder : str | os.PathLike
            Path to simulation results folder.

        Returns
        -------
        str | None
            Returns the code version, except for sphere benchmark, which returns None
        """
        output = self._get_output(simulation_folder)
        try:
            version = output.out.get_code_version()
            return version
        except ValueError:
            logging.warning(
                "Code version not found in the output file or aux file for %s",
                simulation_folder,
            )
            logging.warning(
                "Contents of the directory: %s",
                os.listdir(os.path.dirname(simulation_folder)),
            )
            return None

    def _get_output(self, results_path: str | os.PathLike) -> SphereMCNPSimOutput:
        """
        Method to retrieve output data from MCNP as a SphereMCNPSimOutput

        Parameters
        ----------
        results_path : str | os.PathLike
            Path to simulation results

        Returns
        -------
        output : SphereMCNPSimOutput
            SphereMCNPSimOutput output object
        """
        for file in os.listdir(results_path):
            if file[-1] == "m":
                mfile = file
            elif file[-1] == "o":
                outfile = file

        # Parse output
        mfile = os.path.join(results_path, mfile)
        outfile = os.path.join(results_path, outfile)
        output = SphereMCNPSimOutput(mfile, outfile)
        return output

    def _read_output(self) -> tuple[dict, dict, dict, dict]:
        """Reads all MCNP outputs from a library

        Returns
        -------
        outputs : dic
            Dictionary of MCNP sphere output objects used in plotting, keys are material name or ZAID number
        results : dic
            Dictionary of overview of Tally values for each material/ZAID, returns either all values > 0 for
            tallies with postiive values only, all Values = 0 for empty tallies, and returns the corresponding
            tally bin if it finds any negative values. Contents of the "Values" worksheet.
        errors : dic
            Dictionary of average errors for each tally for each material/Zaid. Contents of the "Errors" worksheet.
        stat_checks : dic
            Dictionary the MCNP statistical check results for each material/ZAID. Contents of the "Statistical
            Checks" Worksheet.
        """
        # Get results
        results = []
        errors = []
        stat_checks = []
        outputs = {}
        # test_path_mcnp = os.path.join(self.test_path, "mcnp")
        for folder in sorted(os.listdir(self.test_path)):
            results_path = os.path.join(self.test_path, folder, "mcnp")
            pieces = folder.split("_")
            # Get zaid
            zaidnum = pieces[-2]
            # Check for material exception
            if zaidnum == "Sphere":
                zaidnum = pieces[-1].upper()
                zaidname = self.mat_settings.loc[zaidnum, "Name"]
            else:
                zaidname = pieces[-1]
            # Get mfile
            for file in os.listdir(results_path):
                if file[-1] == "m":
                    mfile = file
                elif file[-1] == "o":
                    ofile = file
            # Parse output
            output = SphereMCNPSimOutput(
                os.path.join(results_path, mfile), os.path.join(results_path, ofile)
            )

            outputs[zaidnum] = output
            # Adjourn raw Data
            self.raw_data[zaidnum] = output.tallydata
            # Recover statistical checks
            st_ck = output.stat_checks
            # Recover results and precisions
            res, err = output.get_single_excel_data(
                ["2", "4", "6", "12", "14", "24", "34", "22", "32", "44", "46"]
            )
            for dic in [res, err, st_ck]:
                dic["Zaid"] = zaidnum
                dic["Zaid/Mat Name"] = zaidname
            results.append(res)
            errors.append(err)
            stat_checks.append(st_ck)
        return outputs, results, errors, stat_checks

    def _get_output_files(
        self, results_path: str | os.PathLike
    ) -> tuple[str | os.PathLike, str | os.PathLike, str | os.PathLike]:
        """
        Recover the output files from a directory

        Parameters
        ----------
        results_path : str or path
            path where the results are contained.

        Raises
        ------
        FileNotFoundError
            if the required files are not found.

        Returns
        -------
        file1 : str or os.PathLike
            path to the first file
        file2 : str or os.PathLike
            path to the second file
        file3 : str or os.PathLike

        """
        file1 = None
        file2 = None
        file3 = None

        for file_name in os.listdir(results_path):
            if file_name[-1] == "m":
                file1 = file_name
            elif file_name[-1] == "o":
                file2 = file_name
            elif file_name[-4:] == "msht":
                file3 = file_name

        if file1 is None or file2 is None:
            raise FileNotFoundError(
                f"The following path does not contain the required files for {self.code} output: {results_path}"
            )

        file1 = os.path.join(results_path, file1) if file1 else None
        file2 = os.path.join(results_path, file2) if file2 else None
        file3 = os.path.join(results_path, file3) if file3 else None

        return file1, file2, file3


class OpenMCSphereBenchmarkOutput(AbstractSphereBenchmarkOutput):
    def _read_code_version(self, simulation_path: str | os.PathLike) -> str | None:
        """
        Read OpenMC code version from the output file

        Parameters
        ----------
        simulation_path : str | os.PathLike
            Path to OpenMC simulations

        Returns
        -------
        str | None
            OpenMC code version
        """
        _, spfile = self._get_output_files(simulation_path)
        statepoint = omc.OpenMCStatePoint(spfile)
        version = statepoint.version
        return version

    def _get_output(self, results_path: str) -> SphereOpenMCSimOutput:
        """
        Returns SphereOpenMCSimOutput object conating OpenMC simulation output data

        Parameters
        ----------
        results_path : str
            Path to OpenMC simulation ouputs

        Returns
        -------
        output : SphereOpenMCSimOutput
            OpenMC simulation data object
        """
        for file in os.listdir(results_path):
            if "tallies.out" in file:
                outfile = file

        # Parse output
        _, outfile = self._get_output_files(results_path)
        output = SphereOpenMCSimOutput(outfile)
        return output

    def _get_output_files(self, results_path: str | os.PathLike) -> tuple:
        """
        Recover the output files from a directory

        Parameters
        ----------
        results_path : str or path
            path where the results are contained.
        code : str
            code that generated the output ('mcnp' or 'openmc')

        Raises
        ------
        FileNotFoundError
            if the required files are not found.

        Returns
        -------
        file1 : path
            path to the first file
        file2 : path
            path to the second file (only for mcnp)

        """
        file1 = None
        file2 = None

        for file_name in os.listdir(results_path):
            if file_name.endswith(".out"):
                file1 = file_name
            elif file_name.startswith("statepoint"):
                file2 = file_name

        if file1 is None or file2 is None:
            raise FileNotFoundError(
                f"The following path does not contain the required files for {self.code} output: {results_path}"
            )

        file1 = os.path.join(results_path, file1) if file1 else None
        file2 = os.path.join(results_path, file2) if file2 else None

        return file1, file2

    def _read_output(self) -> tuple[dict, dict, dict]:
        """
        Reads all OpenMC outputs from a library

        Returns
        -------
        outputs : dic
            Dictionary of OpenMC sphere output objects used for plotting, keys are material name or ZAID number
        results : dic
            Dictionary of overview of Tally values for each material/ZAID, returns either all values > 0 for
            tallies with postiive values only, all Values = 0 for empty tallies, and returns the corresponding
            tally bin if it finds any negative values. Contents of the "Values" worksheet.
        errors : dic
            Dictionary of average errors for each tally for each material/Zaid. Contents of the "Errors" worksheet.
        """
        # Get results
        results = []
        errors = []
        # stat_checks = []
        outputs = {}
        # test_path_openmc = os.path.join(self.test_path, "openmc")
        for folder in sorted(os.listdir(self.test_path)):
            results_path = os.path.join(self.test_path, folder, "openmc")
            pieces = folder.split("_")
            # Get zaid
            zaidnum = pieces[-2]
            # Check for material exception
            if zaidnum == "Sphere":
                zaidnum = pieces[-1].upper()
                zaidname = self.mat_settings.loc[zaidnum, "Name"]
            else:
                zaidname = pieces[-1]
            # Parse output
            _, outfile = self._get_output_files(results_path)
            output = SphereOpenMCSimOutput(outfile)
            outputs[zaidnum] = output
            # Adjourn raw Data
            self.raw_data[zaidnum] = output.tallydata
            # Recover statistical checks
            # st_ck = output.stat_checks
            # Recover results and precisions
            res, err = output.get_single_excel_data(["4", "14"])
            for dic in [res, err]:
                dic["Zaid"] = zaidnum
                dic["Zaid/Mat Name"] = zaidname
            results.append(res)
            errors.append(err)
            # stat_checks.append(st_ck)
        return (outputs, results, errors, None)  # stat_checks


class SerpentSphereBenchmarkOutput(AbstractSphereBenchmarkOutput):
    def _read_output(self) -> tuple[dict, dict, dict]:
        """
        Reads all Serpent outputs from a library

        NOT YET IMPLEMENTED

        Returns
        -------
        outputs : dic
            Dictionary of Serpent sphere output objects used in plotting, keys are material name or ZAID number
        results : dic
            Dictionary of overview of Tally values for each material/ZAID, returns either all values > 0 for
            tallies with postiive values only, all Values = 0 for empty tallies, and returns the corresponding
            tally bin if it finds any negative values. Contents of the "Values" worksheet.
        errors : dic
            Dictionary of average errors for each tally for each material/Zaid. Contents of the "Errors" worksheet.
        """
        # Get results
        results = []
        errors = []
        stat_checks = []
        outputs = {}
        test_path_serpent = os.path.join(self.test_path, "serpent")
        for folder in os.listdir(test_path_serpent):
            # Call parser here
            continue
        return outputs, results, errors, stat_checks


class SphereTallyOutput:
    def __init__(self) -> None:
        """_summary_

        Raises
        ------
        RuntimeError
            If SphereTallyObject is initialised
        """
        raise RuntimeError("SphereTallyOutput cannot be instantiated")

    def get_single_excel_data(self, tallies2pp: list) -> tuple[dict, dict]:
        """
        Get the excel data of a single MCNP output

        Parameters
        ----------
        tallies2pp : list
            list of tally numbers to post proccess

        Returns
        -------
        tuple[dict, dict]
            _description_
        """
        # TODO this doesn't seem like it will work now...
        data = self.tallydata.set_index(["Energy"])
        results = {}  # Store excel results of different tallies
        errors = {}  # Store average error in different tallies
        heating_res = {}  # Mid-process heating results
        notes = "Negative Bins:"  # Record negative bins here
        initial_notes_length = len(notes)  # To check if notes are registered
        tally_list = [d for _, d in data.groupby(["Tally N."])]
        heating_tallies = ["4", "6", "44", "46"]
        for tally in tally_list:
            tally_num = str(tally["Tally N."].iloc[0])
            tally_description = tally["Tally Description"].iloc[0]
            mean_error = tally["Error"].mean()
            if tally_num in heating_tallies:
                heating_res[tally_num] = tally["Value"].values[0]
            if tally_num in tallies2pp:
                tally_zero = tally[tally["Value"] == 0]
                original_length = len(tally)
                tally = tally[tally["Value"] < 0]
                if len(tally) > 0:
                    res = "Value < 0 in " + str(len(tally)) + " bin(s)"
                    # Get energy bins
                    bins = list(tally.reset_index()["Energy"].values)
                    notes = notes + "\n(" + str(tally_num) + "): "
                    for ebin in bins:
                        notes = notes + str(ebin) + ", "
                    notes = notes[:-2]  # Clear string from excess commas
                elif len(tally_zero) == original_length:
                    res = "Value = 0 for all bins"
                else:
                    res = "Value > 0 for all bins"
                results[tally_description] = res
                errors[tally_description] = mean_error

        if len(heating_res) == 4:
            comp = "Heating comparison [F4 vs F6]"
            try:
                results["Neutron " + comp] = (
                    heating_res["6"] - heating_res["4"]
                ) / heating_res["6"]
            except ZeroDivisionError:
                results["Neutron " + comp] = 0

            try:
                results["Gamma " + comp] = (
                    heating_res["46"] - heating_res["44"]
                ) / heating_res["46"]
            except ZeroDivisionError:
                results["Gamma " + comp] = 0

        # Notes adding
        if len(notes) > initial_notes_length:
            results["Notes"] = notes
        else:
            results["Notes"] = ""

        return results, errors

    def get_comparison_data(self, tallies2pp: list, code: str) -> tuple[list, list]:
        """
        Get Data for single zaid to be used in comparison.

        Parameters
        ----------
        talies2pp : list
            List of tally numbers to postproccess
        code : str
            Either 'mcnp' or 'openmc' to select which tally numbers to use

        Returns
        -------
        results : list
            All results per tally to compare
        columns : list
            Tally names

        """
        # Tallies to post process
        if code == "mcnp":
            binned_tallies = ["12", "22"]
        if code == "openmc":
            binned_tallies = ["4", "14"]
        integral_tallies = []

        # Acquire data
        data = self.tallydata.set_index(["Energy"])
        totalbins = self.totalbin.set_index("Tally Description")
        results = []  # Store data to compare for different tallies
        errors = []
        columns = []  # Tally names and numbers

        # Reorder tallies
        tallies = {}
        tally_list = [d for _, d in data.groupby(["Tally N."])]
        for tally in tally_list:
            tally_str = str(tally["Tally N."].iloc[0])
            if tally_str in tallies2pp:
                tallies[tally_str] = tally
                if tally_str not in binned_tallies:
                    integral_tallies.append(tally_str)

        # Process binned tallies
        for tally_num in binned_tallies:
            tally = tallies[tally_num]
            tally_description = tally["Tally Description"].iloc[0]
            # Get energy bins
            bins = tally.index.tolist()
            for ebin in bins:
                colname = str(ebin) + " [MeV]" + " [t" + tally_num + "]"
                columns.append(colname)
                results.append(tally["Value"].loc[ebin])
                errors.append(tally["Error"].loc[ebin])
            # Add the total bin
            colname = "Total" + " [t" + tally_num + "]"
            columns.append(colname)
            results.append(totalbins["Value"].loc[tally_description])
            errors.append(totalbins["Error"].loc[tally_description])

        # Proccess integral tallies
        for tally_num in integral_tallies:
            tally = tallies[tally_num]
            tally_description = tally["Tally Description"].iloc[0]
            columns.append(tally_description)
            results.append(float(tally["Value"]))
            errors.append(float(tally["Error"]))
        return results, errors, columns


class SphereMCNPSimOutput(MCNPSimOutput, SphereTallyOutput):
    def __init__(self, mfile: str | os.PathLike, outfile: str | os.PathLike) -> None:
        """
        Initialisation function for SphereMCNPSimOutput to create tallydata and totalbin dictionaries.

        Parameters
        ----------
        mfile : str | os.PathLike
            path to mctal file
        outfile : str | os.PathLike
            path to output file
        """
        super().__init__(mfile, outfile)
        self.tallydata, self.totalbin = self._get_tallydata(self.mctal)

    def _get_tallydata(self, mctal: Mctal) -> tuple[pd.DataFrame, pd.DataFrame]:
        """
        Retrieve and organize mctal data. Simplified for sphere leakage case

        Parameters
        ----------
        mctal : Mctal
            F4Eninx Mctal object

        Returns
        -------
        tallydata : pd.DataFrame
            Pandas dataframe containing organised tally data
        totalbin : pd.DataFrame
            Pandas dataframe containing total tally data
        """

        # Extract data
        rows = []
        rowstotal = []
        for t in mctal.tallies:
            num = t.tallyNumber
            des = t.tallyComment[0]
            nCells = t._getNbins("f", False)
            nCora = t._getNbins("i", False)
            nCorb = t._getNbins("j", False)
            nCorc = t._getNbins("k", False)
            nDir = t._getNbins("d", False)
            # usrAxis = t.getAxis("u")
            nUsr = t._getNbins("u", False)
            # segAxis = t.getAxis("s")
            nSeg = t._getNbins("s", False)
            nMul = t._getNbins("m", False)
            # cosAxis = t.getAxis("c")
            nCos = t._getNbins("c", False)
            # ergAxis = t.getAxis("e")
            nErg = t._getNbins("e", False)
            # timAxis = t.getAxis("t")
            nTim = t._getNbins("t", False)

            for f, d, u, s, m, c, e, nt, i, j, k in itertools.product(
                range(nCells),
                range(nDir),
                range(nUsr),
                range(nSeg),
                range(nMul),
                range(nCos),
                range(nErg),
                range(nTim),
                range(nCora),
                range(nCorb),
                range(nCorc),
            ):
                try:
                    erg = t.erg[e]
                except IndexError:
                    erg = None

                val = t._getValue(f, d, u, s, m, c, e, nt, i, j, k, 0)
                err = t._getValue(f, d, u, s, m, c, e, nt, i, j, k, 1)
                if val <= 0:
                    err = np.nan

                row = [num, des, erg, val, err]
                rows.append(row)

            # If Energy binning is involved
            if t.ergTC == "t":
                # 7 steps to get to energy, + 4 for time and mesh directions
                totalbin = t.valsErrors[-1][-1][-1][-1][-1][-1][-1][-1][-1][-1][-1]
                totalvalue = totalbin[0]
                if totalvalue > 0:
                    totalerror = totalbin[-1]
                else:
                    totalerror = np.nan
                row = [num, des, totalvalue, totalerror]
                rowstotal.append(row)

        df = pd.DataFrame(
            rows, columns=["Tally N.", "Tally Description", "Energy", "Value", "Error"]
        )
        dftotal = pd.DataFrame(
            rowstotal, columns=["Tally N.", "Tally Description", "Value", "Error"]
        )
        return df, dftotal


class SphereOpenMCSimOutput(OpenMCSimOutput, SphereTallyOutput):
    def __init__(self, output_path: str | os.PathLike) -> None:
        """
        Initialisation function for SphereOpenMCSimOutput class

        Parameters
        ----------
        output_path : str | os.PathLike
            Path to OpenC simulation output files

        Returns
        -------
        None
        """
        self.output = omc.OpenMCSphereStatePoint(output_path)
        self.tallydata, self.totalbin = self.process_tally()
        self.stat_checks = None

    def _create_dataframe(self, rows: list) -> tuple[pd.DataFrame, pd.DataFrame]:
        """
        Creates dataframe from the data in each output passed through as
        a list of lists from the process_tally function

        Parameters
        ----------
        rows : list
            list of list containing the rows of information from an output file

        Returns
        -------
        df : pd.DataFrame
            dataframe containing the information from each output
        dftotal : pd.DataFrame
            dataframe containing the sum of all values and errors for each output
        """

        df = pd.DataFrame(
            rows, columns=["Tally N.", "Tally Description", "Energy", "Value", "Error"]
        )
        # rowstotal = [[rows[-1][0], rows[-1][1], 0.0, 0.0]]
        # for row in rows:
        #    rowstotal[0][2] += row[3]
        #    rowstotal[0][3] += row[4] ** 2
        # rowstotal[0][3] = np.sqrt(rowstotal[0][3])
        tallies = df["Tally N."].unique().tolist()
        descriptions = df["Tally Description"].unique().tolist()
        rowstotal = []
        for tally, description in zip(tallies, descriptions):
            value = df.loc[df["Tally N."] == tally, "Value"].sum()
            error = np.sqrt((df.loc[df["Tally N."] == tally, "Error"] ** 2).sum())
            rowstotal.append([tally, description, value, error])
        dftotal = pd.DataFrame(
            rowstotal, columns=["Tally N.", "Tally Description", "Value", "Error"]
        )
        return df, dftotal

    def process_tally(self) -> tuple[pd.DataFrame, pd.DataFrame]:
        """
        Creates dataframe from the data in each output passed through as
        a list of lists from the process_tally function

        Returns
        -------
        df : pd.DataFrame
            dataframe containing the information from each output
        dftotal : pd.DataFrame
            dataframe containing the sum of all values and errors for each output
        """
        rows = self.output.tally_to_rows()
        tallydata, totalbin = self._create_dataframe(rows)
        return tallydata, totalbin


class SphereSDDROutput(MCNPSphereBenchmarkOutput):
    times = ["0s", "2.7h", "24h", "11.6d", "30d", "10y"]
    timecols = {
        "0s": "1.0",
        "2.7h": "2.0",
        "24h": "3.0",
        "11.6d": "4.0",
        "30d": "5.0",
        "10y": "6.0",
    }

    def pp_excel_single(self) -> None:
        """
        Generate the single library results excel

        Returns
        -------
        None.

        """
        self.outputs = {}
        self.results = {}
        self.errors = {}
        self.stat_checks = {}
        if self.d1s:
            # template = os.path.join(os.getcwd(), "templates", "SphereSDDR_single.xlsx")
            outpath = os.path.join(
                self.excel_path, "SphereSDDR_single_" + self.lib + ".xlsx"
            )
            # compute the results
            outputs, results, errors, stat_checks = self._compute_single_results()
            self.outputs = outputs
            self.results = results
            self.errors = errors
            self.stat_checks = stat_checks
            lib_name = self.session.conf.get_lib_name(self.lib)
            # Write excel
            # ex = SphereExcelOutputSheet(template, outpath)
            # Results
            # ex.insert_df(11, 2, results, 0, header=False)
            # ex.insert_df(11, 2, errors, 1, header=False)
            # ex.insert_df(9, 2, stat_checks, 2, header=True)
            # lib_name = self.session.conf.get_lib_name(self.lib)
            # ex.wb.sheets[0].range("E1").value = lib_name
            # ex.save()
            exsupp.sphere_sddr_single_excel_writer(
                outpath, lib_name, results, errors, stat_checks
            )

    def pp_excel_comparison(self) -> None:
        """
        Generate the excel comparison output

        Returns
        -------
        None.

        """
        # template = os.path.join(os.getcwd(), "templates", "SphereSDDR_comparison.xlsx")
        if self.d1s:
            for reflib, tarlib, name in self.couples:
                outpath = os.path.join(
                    self.excel_path, "Sphere_SDDR_comparison_" + name + ".xlsx"
                )
                final, absdiff, std_dev = self._compute_compare_result(reflib, tarlib)

                # --- Write excel ---
                # Generate the excel
                # ex = SphereExcelOutputSheet(template, outpath)
                # Prepare the copy of the comparison sheet
                # ws_comp = ex.wb.sheets["Comparison"]
                # ws_diff = ex.wb.sheets["Comparison (Abs diff)"]

                # WRITE RESULTS
                # Percentage comparison
                # rangeex = ws_comp.range("B11")
                # rangeex.options(index=True, header=False).value = final
                # ws_comp.range("E1").value = name

                # Absolute difference comparison
                # rangeex = ws_diff.range("B11")
                # rangeex.options(index=True, header=False).value = absdiff

                single_pp_files = []
                # Add single pp sheets
                for lib in [reflib, tarlib]:
                    pp_dir = self.session.state.get_path(
                        "single", [lib, "SphereSDDR", "d1s", "Excel"]
                    )
                    pp_file = os.listdir(pp_dir)[0]
                    single_pp_path = os.path.join(pp_dir, pp_file)
                    single_pp_files.append(single_pp_path)

                exsupp.sphere_sddr_comp_excel_writer(
                    outpath, name, final, absdiff, std_dev, single_pp_files
                )

    def _get_organized_output(self) -> tuple[list, list, list]:
        """
        Simply recover a list of the zaids and libraries involved

        Returns
        -------
        libs : list
            list of libraries
        zaids : list
            list of zaids
        outputs : list
            list of outputs
        """
        zaids = []

        for (zaidnum, mt, lib), outputslib in self.outputs.items():
            zaids.append((zaidnum, mt))

        zaids = list(set(zaids))
        libs = []  # Not used
        outputs = []  # Not used

        return libs, zaids, outputs

    def _generate_single_plots(self) -> None:
        """
        Method to generate single plots

        Returns
        -------
        None
        """
        libs, allzaids, outputs = self._get_organized_output()
        globalname = self.lib
        self._generate_plots(allzaids, globalname)

    def _generate_plots(self, allzaids: list, globalname: str) -> None:
        """
        Generate all the plots requested by the Sphere SDDR benchmark

        Parameters
        ----------
        allzaids : list
            list of all zaids resulting from the union of the results from
            both libraries.
        globalname : str
            name for the output.

        Returns
        -------
        None.

        """
        # Check if self libraries is already a list
        if type(self.lib) != list:
            libraries = [self.lib]
        else:
            libraries = self.lib

        # Initialize atlas
        outpath = os.path.join(self.atlas_path, "tmp")
        if not os.path.exists(outpath):
            os.mkdir(outpath)
        template = os.path.join(self.path_templates, "AtlasTemplate.docx")
        atlas = at.Atlas(template, "Sphere SDDR " + globalname)
        libmanager = self.session.lib_manager

        # ------------- Binned plots of gamma flux ------------
        atlas.doc.add_heading("Photon Flux (32)", level=1)
        fluxquantity = "Photon Flux"
        fluxunit = r"$p/(cm^2\cdot\#_S)$"
        allzaids.sort()
        # --- Binned plots of the gamma flux ---
        for zaidnum, mt in tqdm(allzaids, desc=" Binned flux plots"):
            material = False
            # Get everything for the title of the zaid
            try:
                name, formula = libmanager.get_zaidname(zaidnum)
                args = [zaidnum, name, formula, mt]
                title = "Zaid: {} ({} {}), MT={}".format(*args)
                # For zaids cooldown time does not change anything
                # Keep the multiple times only for materials
                times = [self.times[0]]
                zaidmatname = f"{formula}_{mt}"
            except ValueError:  # A material is passed instead of zaid
                matname = self.mat_settings.loc[zaidnum, "Name"]
                title = zaidnum + " (" + matname + ")"
                times = self.times
                zaidmatname = f"{matname}_all"
                material = True
            atlas.doc.add_heading(title, level=2)

            # --- Plot the parent contributions ---
            if material:
                for lib in libraries:
                    try:  # Zaid could not be common to the libraries
                        outp = self.outputs[zaidnum, mt, lib]
                    except KeyError:
                        # It is ok, simply nothing to plot here since zaid was
                        # not in library
                        continue
                    try:
                        sddr = outp.tallydata[104].set_index("User")
                    except KeyError:
                        continue  # older version were parents were not tracked

                    lib_name = self.session.conf.get_lib_name(lib)
                    atlas.doc.add_heading(
                        "Parent contribution for {}".format(lib_name), level=3
                    )
                    title = "Parent contribution to SDDR - {}".format(lib_name)
                    libdatas = []

                    tot_dose = sddr.groupby("Time").sum()["Value"].values
                    for parentzaid in set(sddr.index):
                        if int(parentzaid) != 0:
                            _, formula_parent = self.session.lib_manager.get_zaidname(
                                str(abs(parentzaid))
                            )
                            y = sddr.loc[parentzaid]["Value"] / tot_dose * 100
                            libdata = {
                                "x": self.times,
                                "y": y,
                                "err": [],
                                "ylabel": formula_parent,
                            }
                            libdatas.append(libdata)

                    outname = "tmp"
                    quantity = "SDDR contribution"
                    unit = "%"
                    xlabel = "Cooldown time"

                    plot = plotter.Plotter(
                        libdatas,
                        title,
                        outpath,
                        outname,
                        quantity,
                        unit,
                        xlabel,
                        self.testname,
                    )
                    img_path = plot._contribution(legend_outside=True)

                    # Insert the image in the atlas
                    atlas.insert_img(img_path)

            for time in times:
                atlas.doc.add_heading("Cooldown time = {}".format(time), level=3)
                title = "Gamma Leakage flux after a {} cooldown".format(time)
                data = []
                for lib in libraries:
                    try:  # Zaid could not be common to the libraries
                        outp = self.outputs[zaidnum, mt, lib]
                    except KeyError:
                        # It is ok, simply nothing to plot here since zaid was
                        # not in library
                        continue
                    # Get the zaid flux
                    tally_data = outp.tallydata[32].set_index("Time")

                    # Select the correct time
                    t = "F" + self.timecols[time]
                    tally_data = tally_data.loc[t]
                    # If for some reason a total survived just kill him
                    tally_data = tally_data[tally_data.Energy != "total"]

                    energy = tally_data["Energy"].values
                    values = tally_data["Value"].values
                    error = tally_data["Error"].values
                    lib_name = self.session.conf.get_lib_name(lib)
                    ylabel = f"{zaidmatname} ({lib_name})"
                    libdata = {"x": energy, "y": values, "err": error, "ylabel": ylabel}
                    data.append(libdata)

                outname = "{}-{}-{}-{}".format(zaidmatname, globalname, 32, t)
                plot = plotter.Plotter(
                    data,
                    title,
                    outpath,
                    outname,
                    fluxquantity,
                    fluxunit,
                    "Energy [MeV]",
                    self.testname,
                )
                outfile = plot.plot("Binned graph")
                atlas.insert_img(outfile)

        # --- Wave plots flux ---
        # Do this block only if libs are more than one
        lim = 35  # limit of zaids to be in a single plot
        # Plot parameters which are not going to change
        quantity = ["Neutron Flux", "Photon Flux", "SDDR"]
        unit = [r"$n/(cm^2\cdot n_S)$", r"$p/(cm^2\cdot n_S)$", "Sv/h"]
        xlabel = "Zaid/Material and MT value"
        if len(libraries) > 1:
            atlas.doc.add_heading("Flux and SDDR ratio plots", level=1)
            # 1) collect zaid-mt couples in libraries and keep only the ones
            #    that appears on the reference + at least one lib
            # Build a df will all possible zaid, mt, lib combination
            if self.d1s:
                allkeys = list(self.outputs.keys())
            else:
                raise NotImplementedError("Only d1s is implemented")
            df = pd.DataFrame(allkeys)
            df.columns = ["zaid", "mt", "lib"]
            df["zaid-mt"] = df["zaid"].astype(str) + "-" + df["mt"].astype(str)
            df.set_index("lib", inplace=True)
            # get the reference zaids
            refzaids = set(df.loc[self.lib[0]]["zaid-mt"].values)
            otherzaids = set(df.drop(self.lib[0])["zaid-mt"].values)
            # Get the final zaid-mt couples to consider
            zaid_couples = []
            mat_couples = []
            for zaidmt in refzaids:
                if zaidmt in otherzaids:
                    zaid, mt = zaidmt.split("-")
                    if zaid[0] in "mM":
                        mat_couples.append((zaid, mt))
                    else:
                        zaid_couples.append((zaid, mt))
            # sort it
            zaid_couples.sort(key=self._sortfunc_zaidMTcouples)
            mat_couples.sort(key=self._sortfunc_zaidMTcouples)

            # # There is going to be a plot for each cooldown time
            # Only one plot necessary for zaids at cd=0

            # 2) Recover/compute the data that needs to be plot for each lib
            data = []
            time = self.times[0]
            for lib in self.lib:
                nfluxs = []
                pfluxs = []
                sddrs = []
                xlabels = []
                ylabel = self.session.conf.get_lib_name(lib)
                for zaid, mt in zaid_couples:
                    # Extract values
                    try:
                        nflux, pflux, sddr = self._extract_data4plots(
                            zaid, mt, lib, time
                        )
                    except KeyError:
                        # it may be that the zaid is not in the library
                        continue
                    # Memorize values
                    nfluxs.append(nflux)
                    pfluxs.append(pflux)
                    sddrs.append(sddr)
                    name, formula = libmanager.get_zaidname(zaid)
                    xlabels.append(formula + " " + mt)

                # Split the data if its length is more then the limit
                datalenght = len(xlabels)
                sets = math.ceil(datalenght / lim)
                last_idx = 0
                idxs = []
                step = int(datalenght / sets)
                for _ in range(sets):
                    newidx = last_idx + step
                    idxs.append((last_idx, newidx))
                    last_idx = newidx

                for j, (start, end) in enumerate(idxs):
                    # build the dic
                    ydata = [nfluxs[start:end], pfluxs[start:end], sddrs[start:end]]
                    xlab = xlabels[start:end]
                    libdata = {"x": xlab, "y": ydata, "err": [], "ylabel": ylabel}
                    # try to append it to the data in the correct index
                    # if the index is not found, then the list still needs
                    # to be initialized
                    try:
                        data[j].append(libdata)
                    except IndexError:
                        data.append([libdata])

            # 3) Compute parameters for the plotter init
            refname = self.session.conf.get_lib_name(self.lib[0])
            for datapiece in data:
                title = "Ratio Vs {} (T0 + {})".format(refname, time)
                outname = "dummy"  # Does not matter if plot is added imm.
                testname = self.testname
                plot = plotter.Plotter(
                    datapiece, title, outpath, outname, quantity, unit, xlabel, testname
                )
                outfile = plot.plot("Waves")
                atlas.insert_img(outfile)

            # --- Single wave plot for each material ---
            atlas.doc.add_heading("Materials ratio plot", level=1)
            xlab = self.times
            quantity = ["Neutron Flux", "Photon Flux", "SDDR"]
            unit = ["", "", ""]
            xlabel = "Cooldown time"
            for material, _ in tqdm(mat_couples, desc=" Materials: "):
                atlas.doc.add_heading(material, level=2)
                data = []
                for lib in self.lib:
                    ylabel = self.session.conf.get_lib_name(lib)
                    nfluxs = []
                    pfluxs = []
                    sddrs = []
                    for time in self.times:
                        try:
                            nflux, pflux, sddr = self._extract_data4plots(
                                material, "All", lib, time
                            )
                        except KeyError:
                            # it may be that the zaid is not in the library
                            continue
                        # Memorize
                        nfluxs.append(nflux)
                        pfluxs.append(pflux)
                        sddrs.append(sddr)

                    # Build the lib data
                    ydata = [nfluxs, pfluxs, sddrs]

                    libdata = {"x": xlab, "y": ydata, "err": [], "ylabel": ylabel}
                    data.append(libdata)

                # Plot
                refname = self.session.conf.get_lib_name(self.lib[0])
                matname = self.mat_settings.loc[material, "Name"]
                title = "Ratio Vs {} ({})".format(refname, matname)
                outname = "dummy"  # Does not matter if plot is added imm.
                testname = self.testname
                plot = plotter.Plotter(
                    data, title, outpath, outname, quantity, unit, xlabel, testname
                )
                outfile = plot.plot("Waves")
                atlas.insert_img(outfile)

        ########
        print(" Building...")
        if self.d1s:
            atlas.save(self.atlas_path)
        # Remove tmp images
        shutil.rmtree(outpath)

    def _extract_data4plots(
        self, zaid: str, mt: str, lib: str, time: float
    ) -> tuple[float, float, float]:
        """
        Method to extract data for plots

        Parameters
        ----------
        zaid : str
            zaid
        mt : str
            mt reaction number
        lib : str
            library
        time : float
            timestep

        Returns
        -------
        nflux : float
            neutron flux
        pflux : float
            proton flux
        sddr : float
            shut down dose rate
        """
        tallies = self.outputs[zaid, mt, lib].tallydata
        # Extract values
        nflux = tallies[12].set_index("Energy")  # .drop("total")
        nflux = nflux.sum().loc["Value"]
        pflux = (
            tallies[22]
            .groupby("Time")
            .sum(numeric_only=True)
            .loc[int(float(self.timecols[time])), "Value"]
        )
        # a simple set_index is not enough as now dose contribution is split
        # by parent in the materials and needs to be summed up
        sddr = tallies[104].groupby("Time").sum(numeric_only=True)
        sddr["Error"] = sddr["abs_error"] / sddr["Value"]
        sddr = sddr.loc["D" + self.timecols[time], "Value"]
        # Memorize values
        return nflux, pflux, sddr

    def _compute_single_results(
        self,
    ) -> tuple[dict[str, dict], pd.DataFrame, pd.DataFrame, pd.DataFrame]:
        """
        Compute the excel single post processing results and memorize them

        Parameters
        ----------

        Returns
        -------
        outputs: dict[str, dict]
            dictionary of the outputs. the first level is the code level
        results : pd.DataFrame
            global excel datataframe of all values.
        errors : pd.DataFrame
            global excel dataframe of all errors.
        stat_checks : pd.DataFrame
            global excel dataframe of all statistical checks.

        """
        # Get results
        # results = []
        # errors = []
        # stat_checks = []
        desc = " Parsing Outputs: "
        if self.d1s:
            # for folder in tqdm(os.listdir(self.test_path), desc=desc):
            outputs, results, errors, stat_checks = self._parserunmcnp(
                self.test_path, self.lib
            )

            self.outputs = outputs

        # Generate DataFrames
        results = pd.concat(results, axis=1).T
        errors = pd.concat(errors, axis=1).T
        stat_checks = pd.DataFrame(stat_checks)

        # Swap Columns and correct zaid sorting
        # results
        for df in [results, errors, stat_checks]:
            self._sort_df(df)  # it is sorted in place
            df.set_index("Parent")

        # self.outputs = outputs

        return outputs, results, errors, stat_checks

    def _compute_compare_result(
        self, reflib: str, tarlib: str
    ) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
        """
        Given a reference lib and a target lib, both absolute and relative
        comparison are computed

        Parameters
        ----------
        reflib : str
            reference library suffix.
        tarlib : str
            target library suffix.

        Returns
        -------
        final : pd.DataFrame
            relative comparison table.
        absdiff : pd.DataFrame
            absolute comparison table.
        std_dev: pd.DataFrame
            comparison in std. dev. from mean table

        """
        # Get results both of the reflib and tarlib
        comp_dfs = []
        error_dfs = []
        lib_dics = []
        code_outputs = {}
        test_paths = [self.test_path[reflib], self.test_path[tarlib]]
        libs = [reflib, tarlib]
        for test_path, lib in zip(test_paths, libs):
            # Extract all the series from the different reactions
            # Collect the data
            outputs, results, errors, _ = self._parserunmcnp(test_path, lib)
            # Build the df and sort
            comp_df = pd.concat(results, axis=1).T
            error_df = pd.concat(errors, axis=1).T
            for df in [comp_df, error_df]:
                self._sort_df(df)
                # They need to be indexed
                df.set_index(["Parent", "Parent Name", "MT"], inplace=True)
                # Add the df to the list
            comp_dfs.append(comp_df)
            error_dfs.append(error_df)
            lib_dics.append(outputs)
        for dic in lib_dics:
            code_outputs.update(dic)
        self.outputs.update(code_outputs)
        # Consider only common zaids
        idx1 = comp_dfs[0].index
        idx2 = comp_dfs[1].index
        newidx = idx1.intersection(idx2)
        # For some reason they arrive here as objects triggering
        # a ZeroDivisionError
        ref = comp_dfs[0].loc[newidx].astype(float)
        tar = comp_dfs[1].loc[newidx].astype(float)
        ref_err = error_dfs[1].loc[newidx].astype(float)

        # Build the final excel data
        absdiff = ref - tar
        final = absdiff / ref
        std_dev = absdiff / (ref_err * ref)

        # If it is zero the CS are equal! (NaN if both zeros)
        for df in [final, absdiff, std_dev]:
            df.replace(np.nan, "Not Available", inplace=True)
            df.replace(float(0), "Identical", inplace=True)
            df.replace(-np.inf, "Reference = 0", inplace=True)
            df.replace(1, "Target = 0", inplace=True)

        return final, absdiff, std_dev

    @staticmethod
    def _sort_df(df: pd.DataFrame) -> None:
        """
        Sorts the values in the passed dataframe by the Parent column,
        then sets 3 index columns

        Parameters
        ----------
        df : pd.DataFrame
            Dataframe containing output data
        """
        df["index"] = pd.to_numeric(df["Parent"].values, errors="coerce")
        df.sort_values("index", inplace=True)
        del df["index"]

        df.set_index(["Parent", "Parent Name", "MT"], inplace=True)
        df.reset_index(inplace=True)

    @staticmethod
    def _sortfunc_zaidMTcouples(
        item: tuple | list,
    ) -> tuple[bool, str | int, str | int]:
        """
        Function to sort zaid couples

        Parameters
        ----------
        item : tuple | list
            list of zaid couples

        Returns
        -------
        tuple[bool, str | int, str | int]
            (flag, zaid, mt)
        """
        try:
            zaid = int(item[0])
        except ValueError:
            zaid = item[0]
        try:
            mt = int(item[1])
        except ValueError:
            mt = item[1]

        if isinstance(zaid, str):
            flag = True
        else:
            flag = False

        return (flag, zaid, mt)

    def _parserunmcnp(
        self, test_path: str | os.PathLike, lib: str
    ) -> tuple[dict, list, list, list]:
        """
        given a MCNP run folder the parsing of the different outputs is
        performed

        Parameters
        ----------
        test_path : path or str
            path to the test.
        lib : str
            library.

        Returns
        -------
        outputs : Dictionary
            MCNP output object
        results : List
            List of results dataframes
        errors : List
            List of errors dataframes
        stat_checks : List
            List of stat checks dataframes


        """
        results = []
        errors = []
        stat_checks = []
        outputs = {}
        for folder in os.listdir(test_path):
            results_path = os.path.join(test_path, folder, "d1s")
            pieces = folder.split("_")
            # Get zaid
            zaidnum = pieces[1]
            # Check for material exception
            try:
                zaidname = self.mat_settings.loc[zaidnum, "Name"]
                mt = "All"
            except KeyError:
                # it is a simple zaid
                zaidname = pieces[2]
                mt = pieces[3]
            # Get mfile
            for file in os.listdir(results_path):
                if file[-1] == "m":
                    mfile = file
                elif file[-1] == "o":
                    ofile = file
                # Parse output
            output = SphereSDDRMCNPOutput(
                os.path.join(results_path, mfile), os.path.join(results_path, ofile)
            )

            outputs[zaidnum, mt, lib] = output
            # Adjourn raw Data
            self.raw_data[zaidnum, mt, lib] = output.tallydata
            # Recover statistical checks
            st_ck = output.stat_checks
            # Recover results and precisions
            res, err = output.get_single_excel_data()

            for series in [res, err, st_ck]:
                series["Parent"] = zaidnum
                series["Parent Name"] = zaidname
                series["MT"] = mt
            results.append(res)
            errors.append(err)
            stat_checks.append(st_ck)

        return outputs, results, errors, stat_checks

    def print_raw(self) -> None:
        """
        Assigns a path and prints the post processing data as a .csv

        Returns
        -------
        None

        """
        for key, data in self.raw_data.items():
            # Follow the same structure of other benchmarks
            for tallynum, df in data.items():
                filename = "{}_{}_{}.csv".format(key[0], key[1], tallynum)
                file = os.path.join(self.raw_path, filename)
                df.to_csv(file, header=True, index=False)

        # add dump of metadata
        metadata_file = os.path.join(self.raw_path, "metadata.json")
        with open(metadata_file, "w", encoding="utf-8") as outfile:
            json.dump(self.metadata, outfile, indent=4)


class SphereSDDRMCNPOutput(SphereMCNPSimOutput):

    def _get_tallydata(self, mctal: Mctal) -> tuple[dict, dict]:
        """_summary_

        Returns
        -------
        self.tallydata : dict
            dictionary of pandas dataframes containing tally data
        self.totalbin : dict
            dictionary of pandas dataframes containing tally total data
        """
        return self.tallydata, self.totalbin

    @staticmethod
    def _drop_total_rows(df: pd.DataFrame) -> None:
        """
        Method to drop total rows fro dataframes

        Parameters
        ----------
        df : pd.DataFrame
            Dataframe to be sorted
        """
        # drop all total rows
        for key in ["User", "Time", "Energy"]:
            try:
                df.drop(df[df[key] == "total"].index, inplace=True)
            except KeyError:
                pass

    def get_single_excel_data(self) -> tuple[pd.Series, pd.Series]:
        """
        Return the data that will be used in the single
        post-processing excel output for a single reaction

        Returns
        -------
        vals : pd.Series
            series reporting the result of a single reaction.
        errors : pd.Series
            series containing the errors associated with the
            reactions.

        """
        # 32 -> fine gamma flux
        # 104 -> Dose rate
        nflux = self.tallydata[12]
        pflux = self.tallydata[32]
        sddr = self.tallydata[104]
        heat = self.tallydata[46]

        # drop the total rows
        for df in [nflux, pflux, sddr, heat]:
            self._drop_total_rows(df)

        # extend sddr to handle parent contributions abs error is needed
        sddr["abs_error"] = sddr["Error"] * sddr["Value"]

        # Differentiate time labels
        pflux["Time"] = "F" + pflux["Time"].astype(str)
        sddr["Time"] = "D" + sddr["Time"].astype(str)
        heat["Time"] = "H" + heat["Time"].astype(str)

        # Get the total values of the flux at different cooling times
        pfluxvals = pflux.groupby("Time").sum(numeric_only=True)["Value"]
        # Get the mean error of the flux at different cooling times
        pfluxerrors = pflux.groupby("Time").mean(numeric_only=True)["Error"]

        # Get the total values of the SDDR at different cooling times
        sddrvals = sddr.groupby("Time").sum(numeric_only=True)["Value"]
        # Get the mean error of the SDDR at different cooling times
        sddrerrors = sddr.groupby("Time").sum()["abs_error"] / sddrvals

        # Get the total Heating at different cooling times
        heatvals = heat.set_index("Time")["Value"]
        # Get the Heating mean error at different cooling times
        heaterrors = heat.set_index("Time")["Error"]

        # Neutron flux binned in energy
        nfluxvals = nflux.set_index("Energy")["Value"]
        # Errors of the neutron flux
        nfluxerrors = nflux.set_index("Energy")["Error"]

        # # Delete the total row in case it is there
        # for df, tag in zip(
        #     [pfluxvals, pfluxerrors, sddrvals, sddrerrors, heatvals, heaterrors],
        #     ["F", "F", "D", "D", "H", "H"],
        # ):
        #     try:
        #         del df[tag + "total"]
        #     except KeyError:
        #         # If total value is not there it is ok
        #         pass

        # # Do the same for the flux
        # for df in [nfluxvals, nfluxerrors]:
        #     try:
        #         del df["total"]
        #     except KeyError:
        #         # If total value is not there it is ok
        #         pass

        # 2 series need to be built here, one for values and one for errors
        vals = pd.concat([pfluxvals, sddrvals, heatvals, nfluxvals], axis=0)
        errors = pd.concat([pfluxerrors, sddrerrors, heaterrors, nfluxerrors], axis=0)

        return vals, errors
