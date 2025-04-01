import io
import logging
import os

# import win32com.client
# import aspose.words
import docx

# from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

# from docx.oxml import OxmlElement, parse_xml
# from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches
from matplotlib.figure import Figure
import matplotlib.pyplot as plt

from jade.helper.aux_functions import PathLike


class Atlas:
    def __init__(self, template: PathLike, name: str):
        """
        Atlas of plots for post-processing

        Parameters
        ----------
        template : PathLike
            path to the word template file.
        name : str
            name of the atlas file.

        Returns
        -------
        None.

        """
        self.name = name  # Name of the Atlas (from libraries)
        # Open The Atlas template
        doc = docx.Document(template)
        doc.add_heading("JADE ATLAS: " + name, level=0)
        self.outname = "atlas_" + name  # Name for the outfile
        self.doc = doc  # Word Document

    def insert_img(self, figure: Figure, width=Inches(7.5)):
        """Insert an image in the word document

        Parameters
        ----------
        figure : Figure
            matplotlib figure to insert.
        width : Inches, optional
            width in docx Inches, by default Inches(7.5)
        """
        # Convert the figure to an in-memory binary stream
        img_stream = io.BytesIO()
        # Be sure to include extra artists if present
        extra_artists = figure.get_default_bbox_extra_artists()
        for ax in figure.get_axes():
            extra_artists.extend(ax.get_children())
        figure.savefig(
            img_stream,
            format="png",
            dpi=200,
            bbox_inches="tight",
            bbox_extra_artists=extra_artists,
        )
        img_stream.seek(0)

        self.doc.add_picture(img_stream, width=width)
        # be sure to close the figure once it has been added
        plt.close(figure)
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # def insert_df(
    #     self,
    #     df,
    #     caption=None,
    #     highlight=False,
    #     tablestyle=None,  # , template_idx=None,
    # ):
    #     """
    #     Inser a dataframe as a table in a Word file

    #     Parameters
    #     ----------
    #     df : pd.DataFrame
    #         dataframe to insert.
    #     caption : str, optional
    #         caption of the table. The default is None
    #     highlight : bool, optional
    #         Very specific for stress assessment. The default is False.
    #     # template_idx : int, optional
    #     #     index of the template table to use. The default is None
    #     tablestyle : str, optional
    #         table word style to apply. The default is None

    #     Returns
    #     -------
    #     table : docx.Table
    #         table inserted.

    #     """
    #     # Create the table or inherit a template
    #     template_idx = None  # other possibilities not implemented anymore
    #     if template_idx is None:
    #         table = self.doc.add_table(1, len(df.columns))
    #     else:
    #         template = self.template_tables[template_idx]
    #         table = template.insert(self.doc)

    #     # Assign style if provided
    #     if table.style is not None:
    #         table.style = tablestyle

    #     # If template is not provided, the header row must be filled
    #     if template_idx is None:
    #         # add the header rows.
    #         for j in range(df.shape[-1]):
    #             table.cell(0, j).text = df.columns[j]

    #     # Add the rest of the data frame
    #     # val = df.values
    #     for i, (idx, row) in enumerate(df.iterrows()):
    #         # Understand is safety margin is barely acceptable
    #         flag_almost = False
    #         try:
    #             sm = float(row["Safety Margin"])
    #             if sm > 1 and sm < 1.1:
    #                 flag_almost = True
    #         except KeyError:
    #             pass
    #         except ValueError:
    #             pass
    #         except TypeError:
    #             # cannot convert to float
    #             pass

    #         row_cells = table.add_row().cells
    #         for j, item in enumerate(row):
    #             cell = row_cells[j]
    #             cell.text = str(item)
    #             cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #             for par in cell.paragraphs:
    #                 par.style = "Table"
    #             if highlight is not None:
    #                 if cell.text == "NOK":
    #                     self._highlightCell(cell)
    #                 elif cell.text == "OK" and flag_almost:
    #                     self._highlightCell(cell, color="FFFF46")

    #     if caption is not None:
    #         paragraph = self.doc.add_paragraph("Table ", style="Didascalia")
    #         self._wrapper(paragraph, "table")
    #         paragraph.add_run(" - " + caption)
    #     # paragraph = doc.add_paragraph('Figure Text', style='Didascalia')

    #     return table

    def save(self, outpath, pdfprint=True):
        """
        Save word atlas and possibly export PDF

        Parameters
        ----------
        outpath : str/path
            path to export the atlas
        pdfprint : Boolean, optional
            If True export also in PDF format

        Returns
        -------
        None.

        """
        outpath_word = os.path.join(outpath, self.outname + ".docx")
        # outpath_pdf = os.path.join(outpath, self.outname + ".pdf")
        if len(outpath_word) > 259:
            logging.warning(
                "The path to the word document is too long, the file will be truncated"
            )
            outpath_word = outpath_word[:254] + ".docx"

        try:
            self.doc.save(outpath_word)
        except FileNotFoundError as e:
            # If there is still an error it may be due to special char
            # print the original exception
            print(" The following is the original exception:")
            print(e)
            print(
                "\n it may be due to invalid characters in the file name or a path too long"
            )

        # Remove PDF printing. If required, word document can be saved manually.
        if pdfprint:
            pass

    # @staticmethod
    # def _wrapper(paragraph, ptype):
    #     """
    #     Wrap a paragraph in order to add cross reference

    #     Parameters
    #     ----------
    #     paragraph : docx.Paragraph
    #         image to wrap.
    #     ptype : str
    #         type of paragraph to wrap

    #     Returns
    #     -------
    #     None.

    #     """
    #     if ptype == "table":
    #         instruction = " SEQ Table \\* ARABIC"
    #     elif ptype == "figure":
    #         instruction = " SEQ Figure \\* ARABIC"
    #     else:
    #         raise ValueError(ptype + " is not a supported paragraph type")

    #     run = run = paragraph.add_run()
    #     r = run._r
    #     fldChar = OxmlElement("w:fldChar")
    #     fldChar.set(qn("w:fldCharType"), "begin")
    #     r.append(fldChar)
    #     instrText = OxmlElement("w:instrText")
    #     instrText.text = instruction
    #     r.append(instrText)
    #     fldChar = OxmlElement("w:fldChar")
    #     fldChar.set(qn("w:fldCharType"), "end")
    #     r.append(fldChar)

    # @staticmethod
    # def _highlightCell(cell, color="FBD4B4"):
    #     shading_elm_1 = parse_xml(
    #         r'<w:shd {} w:fill="'.format(nsdecls("w")) + color + r'"/>'
    #     )
    #     cell._tc.get_or_add_tcPr().append(shading_elm_1)
