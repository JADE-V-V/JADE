# -*- coding: utf-8 -*-
# Created on Wed Oct 21 17:18:07 2020

# @author: Davide Laghi

# Copyright 2021, the JADE Development Team. All rights reserved.

# This file is part of JADE.

# JADE is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# JADE is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with JADE.  If not, see <http://www.gnu.org/licenses/>.

from __future__ import annotations

import json
import math
import os
import re
import shutil
from abc import abstractmethod

import numpy as np
import pandas as pd
from docx.shared import Inches
from f4enix.input.MCNPinput import D1S_Input
from scipy.interpolate import interp1d
from tqdm import tqdm

import jade.atlas as at
from jade.output import MCNPBenchmarkOutput, MCNPSimOutput
from jade.plotter import Plotter
from jade.status import EXP_TAG

FOILS_REACTION = {
    "Al": "Al(n,alpha)",
    "S": "S-32(n,p)",
    "In": "In-115(n,n')",
    "Rh": "Rh-103(n,n')",
    "Au": "Au-197(n,gamma)",
}

MCNP_UNITS = {"Energy": "MeV", "Time": "shakes"}

TALLY_NORMALIZATION = {
    "Tiara-BC": "lethargy",
    "FNS-TOF": "lethargy",
    "Oktavian": "lethargy",
    "TUD-Fe": "energy bins",
    "TUD-W": "energy bins",
    "TUD-FNG": "energy bins",
    "ASPIS-PCA-Replica_flux": "lethargy",
}

ACTIVATION_REACTION = {
    "Ni-n2n": "Ni-58(n,2n)Ni-57",
    "Al": "Al-27(n,a)Na-24",
    "Fe": "Fe-56(n,p)Mn-56",
    "Ni-np": "Ni-58(n,p)Co-58",
    "Nb": "Nb-93(n,2n)Nb-92",
    "In": "In-115(n,n')In-115m",
    "Mn": "Mn-55(n,g)Mn-56",
    "Au": "Au-197(n,g)Au-198",
    "Rh": "Rh-103(n,n')Rh-103*",
    "S": "S-32(n,p)P-32",
    "Zr": "Zr-90(n,2n)Zr-89",
}


class ExperimentalOutput(MCNPBenchmarkOutput):
    def __init__(self, *args, **kwargs) -> None:
        """
        This extends the Benchmark Output and creates an abstract class
        for all experimental outputs.
        Parameters
        ----------
        *args : TYPE
            see BenchmarkOutput doc.
        **kwargs : TYPE
            see BenchmarkOutput doc.
        multiplerun : bool
            this additional keyword specifies if the benchmark is composed
            by more than one MCNP run. It defaults to False.
        Returns
        -------
        None.

        """
        # Add a special keyword for experimental benchmarks
        try:
            multiplerun = kwargs.pop("multiplerun")
        except KeyError:
            # Default to False
            multiplerun = False
        # Recover session and testname
        session = args[3]
        testname = args[2]
        super().__init__(*args, **kwargs)
        # The experimental data needs to be loaded
        self.path_exp_res = os.path.join(session.path_exp_res, testname)

        # Add the raw path data (not created because it is a comparison)
        out = os.path.dirname(self.atlas_path)
        raw_path = os.path.join(out, "Raw_Data")
        if not os.path.exists(raw_path):
            os.mkdir(raw_path)
        self.raw_path = raw_path
        self.multiplerun = multiplerun

        # Read the metadata from the simulations
        metadata = {}
        for lib, test_path in self.test_path.items():
            if lib == EXP_TAG:
                continue
            code = args[1]
            if self.multiplerun:
                # I still need only one metadata. They should be all the same
                results_path = os.path.join(test_path, os.listdir(test_path)[0], code)
                metadata_lib = self._read_metadata_run(results_path)
            else:
                results_path = os.path.join(test_path, code)
                self.metadata_lib = self._read_metadata_run(results_path)
            metadata[lib] = metadata_lib
        self.metadata = metadata

    def single_postprocess(self) -> None:
        """
        Always raise an Attribute Error since no single post-processing is
        foreseen for experimental benchmarks
        Raises
        ------
        AttributeError
            DESCRIPTION.
        Returns
        -------
        None.
        """
        raise AttributeError("\n No single pp is foreseen for exp benchmark")

    def compare(self) -> None:
        """
        Complete the routines that perform the comparison of one or more
        libraries results with the experimental ones.
        Returns
        -------
        None.
        """
        print(" Exctracting outputs...")
        self._extract_outputs()

        print(" Read experimental results....")
        self._read_exp_results()

        print(" Dumping raw data...")
        self._print_raw()

        print(" Generating Excel Recap...")
        self.pp_excel_comparison()

        print(" Creating Atlas...")
        self.build_atlas()

    def pp_excel_comparison(self) -> None:
        """
        At the moment everything is handled by _pp_excel_comparison that needs
        to be implemented in each child class. Some standard procedures may be
        added in the feature in order to reduce the amount of ex-novo coding
        necessary to implement a new experimental benchmark.
        Returns
        -------
        None.
        """
        self._pp_excel_comparison()

    def build_atlas(self) -> None:
        """
        Creation and saving of the atlas are handled by this function while
        the actual filling of the atlas is left to _build_atlas which needs
        to be implemented for each child class.
        Returns
        -------
        None.
        """
        # Build a temporary folder for images
        tmp_path = os.path.join(self.atlas_path, "tmp")

        os.mkdir(tmp_path)

        globalname = ""
        for lib in self.lib:
            globalname = globalname + lib + "_Vs_"
        globalname = globalname[:-4]
        globalname = self.testname + "_" + globalname
        # Initialize the atlas
        template = os.path.join(self.session.path_templates, "AtlasTemplate.docx")
        atlas = at.Atlas(template, globalname)

        # Fill the atlas
        atlas = self._build_atlas(tmp_path, atlas)

        atlas.save(self.atlas_path)
        # Remove tmp images
        shutil.rmtree(tmp_path)

    def _extract_single_output(
        self, results_path: str | os.PathLike, folder: str, lib: str
    ) -> tuple[pd.DataFrame, str]:
        """Method to extract single output data from MCNP files

        Parameters
        ----------
        results_path : str | os.PathLike
            Path to simulations results.
        folder : str
            Sub-folder for multiple run case. 
        lib : str
            Test library.

        Returns
        -------
        tallydata : pd.DataFrame
            Pandas dataframe containing tally data.
        input : str
            Test name.
        """
        mfile, ofile, meshtalfile = self._get_output_files(results_path)
        # Parse output
        output = MCNPSimOutput(mfile, ofile, meshtalfile)

        # need to extract the input in case of multi
        if self.multiplerun:
            pieces = folder.split("_")
            input = pieces[-1]
            if input not in self.inputs:
                self.inputs.append(input)
            self.outputs[input, lib] = output
            # Get the meaningful results
            self.results[input, lib] = self._processMCNPdata(output)
        else:
            # just treat it as a special case of multiple run
            self.outputs[self.testname, lib] = output
            # Get the meaningful results
            self.results[self.testname, lib] = self._processMCNPdata(output)
            input = self.testname

        return output.tallydata, input

    def _extract_outputs(self) -> None:
        """
        Extract, organize and store the results coming from the different codes
        runs

        Returns
        -------
        None.
        """
        self.outputs = {}
        self.results = {}

        # Each output object is processing only one code at the time at the moment
        if self.mcnp:
            code_tag = "mcnp"
        if self.openmc:
            print("Experimental comparison not implemented for OpenMC")
            return
        if self.serpent:
            print("Experimental comparison not implemented for Serpent")
            return
        if self.d1s:
            code_tag = "d1s"

        # only multiple runs have multiple inputs
        if self.multiplerun:
            self.inputs = []
        else:
            self.inputs = [self.testname]

        # Iterate on the different libraries results except 'Exp'
        for lib, test_path in self.test_path.items():
            if lib != EXP_TAG:
                if self.multiplerun:
                    # Results are organized by folder and lib
                    code_raw_data = {}
                    for folder in os.listdir(test_path):
                        results_path = os.path.join(test_path, folder, code_tag)
                        tallydata, input = self._extract_single_output(
                            results_path, folder, lib
                        )
                        code_raw_data[input, lib] = tallydata

                # Results are organized just by lib
                else:
                    results_path = os.path.join(test_path, code_tag)
                    tallydata, input = self._extract_single_output(
                        results_path, None, lib
                    )
                    code_raw_data = {(self.testname, lib): tallydata}

                # Adjourn raw Data
                # self.raw_data[code_tag].update(code_raw_data)
                self.raw_data.update(code_raw_data)

    def _read_exp_results(self) -> None:
        """
        Read all experimental results and organize it in the self.exp_results
        dictionary.
        If multirun is set to true the first layer of the dictionary will
        consist in the different folders and the second layer will be the
        different files. If it is not multirun, insetead, only one layer of the
        different files will be generated.
        All files need to be in .csv format. If a more complex format is
        provided, the user should ovveride the _read_exp_file method.
        Returns
        -------
        None.
        """
        exp_results = {}
        if self.multiplerun:
            # Iterate on each folder and then in each file, read them and
            # build the result dic
            for folder in os.listdir(self.path_exp_res):
                exp_results[folder] = {}
                cp = os.path.join(self.path_exp_res, folder)
                for file in os.listdir(cp):
                    filename = file.split(".")[0]
                    filepath = os.path.join(cp, file)
                    df = self._read_exp_file(filepath)
                    c = df.columns.tolist()[1]
                    df = df[df[c] > 2e-38]
                    exp_results[folder][filename] = df
        else:
            # Iterate on each each file, read it and
            # build the result dic
            exp_results[self.testname] = {}
            for file in os.listdir(self.path_exp_res):
                filename = file.split(".")[0]
                filepath = os.path.join(self.path_exp_res, file)
                df = self._read_exp_file(filepath)
                c = df.columns.tolist()[1]
                df = df[df[c] > 2e-38]
                exp_results[self.testname][filename] = df

        self.exp_results = exp_results

    @staticmethod
    def _read_exp_file(filepath : str | os.PathLike) -> pd.DataFrame:
        """
        Default way of reading a csv file
        Parameters
        ----------
        filepath : path/str
            experimental file results to be read.
        Returns
        -------
        pd.DataFrame
            Contain the data read.
        """
        return pd.read_csv(filepath)

    def _print_raw(self) -> None:
        """
        Dump all the raw data
        Returns
        -------
        None.
        """
        raw_to_print = self.raw_data.items()

        for (folder, lib), item in raw_to_print:
            # Create the lib directory if it is not there
            cd_lib = os.path.join(self.raw_path, lib)
            if not os.path.exists(cd_lib):
                os.mkdir(cd_lib)
                # dump also the metadata if it is the first time
                with open(
                    os.path.join(cd_lib, "metadata.json"), "w", encoding="utf-8"
                ) as f:
                    json.dump(self.metadata[lib], f)

            # Dump everything
            for key, data in item.items():
                if folder == self.testname:
                    file = os.path.join(cd_lib, str(key) + ".csv")
                else:
                    file = os.path.join(cd_lib, folder + " " + str(key) + ".csv")
                data.to_csv(file, header=True, index=False)

    @abstractmethod
    def _processMCNPdata(self, output : MCNPSimOutput) -> dict:
        """
        Given an mctal file object return the meaningful data extracted. Some
        post-processing on the data may be foreseen at this stage.
        Parameters
        ----------
        output : MCNPoutput
            object representing an MCNP output.
        Returns
        -------
        item :
            the type of item can vary based on what the user intends to do
            whith it. It will be stored in an organized way in the self.results
            dictionary
        """
        item = None
        return item

    @abstractmethod
    def _pp_excel_comparison(self) -> None:
        """
        Responsible for producing excel outputs
        Returns
        -------
        """
        pass

    @abstractmethod
    def _build_atlas(self, tmp_path : str | os.PathLike, atlas : at.Atlas) -> at.Atlas:
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.
        Parameters
        ----------
        tmp_path : path
            path to the temporary folder where to dump images.
        atlas : Atlas
            Object representing the plot Atlas.
        Returns
        -------
        atlas : Atlas
            After being filled the atlas is returned.
        """
        atlas = None
        return atlas


class FNGOutput(ExperimentalOutput):
    names = ["FNG1", "FNG2"]
    times = {
        "FNG1": ["1d", "7d", "15d", "30d", "60d"],
        "FNG2": [
            "1.22h",
            "1.72h",
            "2.08h",
            "3.22h",
            "4.80h",
            "6.80h",
            "9.47h",
            "12.7h",
            "15.9h",
            "20.2h",
            "25.2h",
            "1.53d",
            "2.46d",
            "4d",
            "5.55d",
            "8.20d",
            "12.2d",
            "19.3d",
            "19.8d",
        ],
    }

    def _processMCNPdata(self, output: MCNPSimOutput):
        """
        Read All tallies and return them as a dictionary of DataFrames. This
        aslo needs to ovveride the raw data since unfortunately it appears
        that the user bins necessary for tracking daughters and parents are
        not correclty written to the mctal file.
        Parameters
        ----------
        output : MCNPoutput
            object representing the MCNP output.
        Returns
        -------
        df : pd.DataFrame
            table of the SDDR at different cooling timesbadi
        """
        res = {}
        mctal = output.mctal
        # Cutom of read of tallies due to errors in the mctal file
        for tally in mctal.tallies:
            tallyres = []
            tnum = int(tally.tallyNumber)

            # -- Get SDDR --
            if tnum == 4:
                for i, time in enumerate(tally.tim):
                    val = tally._getValue(0, 0, 0, 0, 0, 0, 0, i, 0, 0, 0, 0)
                    err = tally._getValue(0, 0, 0, 0, 0, 0, 0, i, 0, 0, 0, 1)

                    # Store
                    time_res = [i + 1, val, err]
                    tallyres.append(time_res)

                # Build and store the taly df
                df = pd.DataFrame(tallyres)
                df.columns = ["time", "sddr", "err"]
                res[str(tnum)] = df

            # -- Parent tracker --
            if tnum in [14, 24]:
                for i in range(tally.nTim):
                    for j in range(tally.nUsr):
                        val = tally._getValue(0, 0, j, 0, 0, 0, 0, i, 0, 0, 0, 0)
                        err = tally._getValue(0, 0, j, 0, 0, 0, 0, i, 0, 0, 0, 1)
                        # Store
                        time_res = [i + 1, j, val, err]
                        tallyres.append(time_res)

                # Build and store the taly df
                df = pd.DataFrame(tallyres)
                df.columns = ["time", "tracked", "sddr", "err"]
                # The first row is the complementary bin (0) and last row
                # is the total. They can be dropped
                df = df.set_index("tracked").drop([0, j]).reset_index()
                res[str(tnum)] = df

        # --- Override the raw data ---
        # Get the folder and lib
        path = output.mctal_file
        folderpath = os.path.dirname(os.path.dirname(path))
        folder = os.path.basename(folderpath)
        lib = os.path.basename(os.path.dirname(os.path.dirname(folderpath)))
        self.raw_data[folder, lib] = res

        return res

    def _pp_excel_comparison(self) -> None:
        """
        Responsible for producing excel outputs
        """
        # Dump the global C/E table
        ex_outpath = os.path.join(self.excel_path, self.testname + "_CE_tables.xlsx")
        # Create a Pandas Excel writer using XlsxWriter as the engine.
        with pd.ExcelWriter(ex_outpath, engine="xlsxwriter") as writer:
            # --- build and dump the C/E table ---
            for folder in self.names:
                # collect all available data
                alldata = self._get_collected_data(folder)
                exp_err = alldata["Relative Error"]
                exp_sddr = alldata["Experimental SDDR [Sv/h]"]
                # build the C\E df
                df = pd.DataFrame(index=alldata.index)

                for lib in self.lib[1:]:
                    libname = self.session.conf.get_lib_name(lib)
                    # get computational data
                    com_err = alldata[lib + "err"]
                    com_sddr = alldata[lib + "sddr"]

                    # compute global error (SRSS)
                    gl_err = ((com_err**2 + exp_err**2) ** (1 / 2)).round(2).astype(str)
                    # compute C/E
                    gl_val = (com_sddr / exp_sddr).round(2).astype(str)

                    df[libname] = gl_val + " +/- " + gl_err

                # Dump the df
                df.to_excel(writer, sheet_name=folder, startrow=2)
                # Write description
                ws = writer.sheets[folder]
                ws.write_string(0, 0, '"C/E (mean +/- σ)"')

    def _get_collected_data(self, folder):
        """
        Given a campaign it builds a single table containing all experimental
        and computational data available for the total SDDR tally.
        Parameters
        ----------
        folder : str
            campaign name.
        Returns
        -------
        df : pd.DataFrame
            collective data on the campaing.
        """
        idx = ["Cooldown Time [s]", "Cooldown Time [d]"]
        # Initialize the table with the experimental results
        df = self.exp_results[folder]["SDDR"].copy()
        df = df.set_index(idx).sort_index()

        # Avoid exp tag
        for lib in self.lib[1:]:
            libdf = self.results[folder, lib]["4"].set_index("time").sort_index()
            # add the SDDR and relative column of each library
            df[lib + "sddr"] = libdf["sddr"].values
            df[lib + "err"] = libdf["err"].values

        return df

    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.
        Parameters
        ----------
        tmp_path : path
            path to the temporary folder where to dump images.
        atlas : Atlas
            Object representing the plot Atlas.
        Returns
        -------
        atlas : Atlas
            After being filled the atlas is returned.
        """
        patzaid = re.compile(r"(?<=[\s\-\t])\d+(?=[\s\t\n])")

        atlas.doc.add_heading("Shut Down Dose Rate", level=1)
        xlabel = "Cooldown time"
        # Only two plots, one for each irradiation campaign
        for folder, title in zip(
            self.names, ["1st FNG Irradiation campaign", "2nd FNG Irradiation campaign"]
        ):
            atlas.doc.add_heading(title, level=2)
            # --- SDDR PLOT ---
            # -- Recover data to plot --
            data = []
            x = self.times[folder]
            for lib in self.lib:
                if lib == "Exp":
                    df = self.exp_results[folder]["SDDR"]
                    y = df["Experimental SDDR [Sv/h]"].values
                    err = (df["Relative Error"] * y).values
                    ylabel = "Experiment"
                else:
                    df = self.results[folder, lib]["4"].set_index("time").sort_index()
                    y = df.sddr.values
                    err = df.err.values * y
                    ylabel = self.session.conf.get_lib_name(lib)

                data.append({"x": x, "y": y, "err": err, "ylabel": ylabel})
            # -- Plot --
            outname = "tmp"
            quantity = "SDDR"
            unit = "Sv/h"
            plot = Plotter(
                data, title, tmp_path, outname, quantity, unit, xlabel, self.testname
            )
            img_path = plot.plot("Discrete Experimental points")
            # Insert the image in the atlas
            atlas.insert_img(img_path)

            # --- Tracking PLOTs ---
            # -- Recover data to plot --
            # There is the need to recover the tracked parents and daughters
            zaid_tracked = {}
            for lib in self.lib[1:]:
                file = os.path.join(self.test_path[lib], folder, "d1s", folder)
                inp = D1S_Input.from_input(file)
                for tallynum in ["24", "14"]:
                    card = inp.get_data_cards("FU" + tallynum)["FU" + tallynum]
                    strings = []
                    for line in card.lines:
                        zaids = patzaid.findall(line)
                        for zaid in zaids:
                            if zaid != "0":
                                _, formula = self.session.lib_manager.get_zaidname(zaid)
                                strings.append(formula)

                    zaid_tracked[tallynum] = strings

            x = self.times[folder]
            titles = {
                "parent": title + ", parent isotopes contribution ",
                "daughter": title + ", daughter isotopes contribution ",
            }
            tallynums = {"parent": "24", "daughter": "14"}

            for tracked in ["parent", "daughter"]:
                atlas.doc.add_heading(tracked + " tracking", level=3)
                for lib in self.lib[1:]:
                    libname = self.session.conf.get_lib_name(lib)

                    # Recover the data
                    tallynum = tallynums[tracked]
                    df = self.results[folder, lib][tallynum]
                    zaidstracked = set(df.tracked.values)
                    tot_dose = df.groupby("time").sum().sddr.values
                    df.set_index("tracked", inplace=True)
                    data = []
                    for i, zaid in enumerate(zaidstracked):
                        subset = df.loc[zaid]
                        assert len(subset.time.values) == len(x)
                        formula = zaid_tracked[tallynum][i]
                        y = subset.sddr.values / tot_dose * 100
                        libdata = {"x": x, "y": y, "err": [], "ylabel": formula}
                        data.append(libdata)

                    outname = "tmp"
                    newtitle = titles[tracked] + libname
                    quantity = "SDDR contribution"
                    unit = "%"
                    xlabel = "Cooldown time"

                    plot = Plotter(
                        data,
                        newtitle,
                        tmp_path,
                        outname,
                        quantity,
                        unit,
                        xlabel,
                        self.testname,
                    )
                    img_path = plot._contribution(legend_outside=True)

                    # Insert the image in the atlas
                    atlas.insert_img(img_path)

        return atlas

    def _read_exp_file(self, filepath):
        """
        Override parent method since the separator for these experimental
        files is ";"
        Parameters
        ----------
        filepath : str
            string containing the path to the experimental file to be read
            for comparison

        """
        return pd.read_csv(filepath, sep=";")


class SpectrumOutput(ExperimentalOutput):
    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str
            path to the temporary folder containing the plots for the atlas
        atlas : Atlas
            Object representing the plot Atlas.

        """
        self.tables = []
        self.bench_conf = pd.read_excel(self.cnf_path)
        self.bench_conf = self.bench_conf.set_index(["Tally"])
        # Loop over benchmark cases
        for input in tqdm(self.inputs, desc=" Inputs: "):
            # Loop over tallies
            for tally in self.outputs[(input, self.lib[1])].mctal.tallies:
                # Get tally number and info
                tallynum, particle, xlabel = self._get_tally_info(tally)
                # Collect data
                quantity_CE = self.bench_conf.loc[tallynum, "Y Label"]
                e_int = self.bench_conf.loc[tallynum, "C/E X Quantity intervals"]
                e_int = e_int.split("-")

                # Convert the list of number strings into a list of integers
                e_intervals = [float(num) for num in e_int]
                data, xlabel = self._data_collect(
                    input, str(tallynum), quantity_CE, e_intervals
                )
                if not data:
                    continue

                # Use re.findall to extract all substrings between '[' and ']'
                unit = self.bench_conf.loc[tallynum, "Y Unit"]
                quantity = self.bench_conf.loc[tallynum, "Quantity"]
                title = self._define_title(input, quantity_CE)
                atlas.doc.add_heading(title, level=1)
                # Once the data is collected it is passed to the plotter
                outname = "tmp"
                plot = Plotter(
                    data,
                    title,
                    tmp_path,
                    outname,
                    quantity,
                    unit,
                    xlabel,
                    self.testname,
                )
                img_path = plot.plot("Experimental points")
                # Insert the image in the atlas
                atlas.insert_img(img_path)

        # Dump C/E table
        self._dump_ce_table()

        return atlas

    def _get_tally_info(self, tally):
        """
        Extracts and assigns information from the tally object, as well as
        information from the benchmark config variable

        Args:
            tally (Tally): JADE tally object

        Returns:
            tallynum (int): Tally number of the tally being plotted
            particle (str): Type of quantity being plotted on the X axis
            quant + unit (str): Unit of quantity being plotted on the X axis

        """
        tallynum = tally.tallyNumber
        particle = tally.particleList[np.where(tally.tallyParticles == 1)[0][0]]
        quant = self.bench_conf.loc[tallynum, "X Quantity"]
        unit = self.bench_conf.loc[tallynum, "X Unit"]
        return tallynum, particle, quant + " [" + unit + "]"

    def _define_title(self, input, quantity_CE):
        """Assigns the title for atlas plot

        Parameters
        ----------
        input : str
            String containing the name of the benchmark case being run, will
            be subfolder test name for benchmarks with multiple runs
        quantity_CE : str
            String containing the Y axis variable from benchmark config file

        Returns
        -------
        title : str
            Title string
        """

        if not self.multiplerun:
            title = self.testname + ", " + quantity_CE
        else:
            title = self.testname + " " + input + ", " + quantity_CE
        return title

    def _dump_ce_table(self):
        """
        Generates the C/E table and dumps them as an .xlsx file
        """
        final_table = pd.concat(self.tables)
        skipcol_global = 0
        binning_list = ["Energy", "Time"]
        for x_ax in binning_list:  # to update if other binning will be used
            x_lab = x_ax[0]
            col_check = "Max " + x_lab
            ft = final_table.set_index(["Input"])

            if col_check not in final_table.columns.tolist():
                continue
            else:
                todump = final_table.set_index(["Input", "Quantity", "Library"])
            for binning in binning_list:
                if binning == x_ax:
                    continue
                else:
                    # if tallies only have one type of binning KeyError could
                    # arise
                    try:
                        todump = todump.drop(
                            columns=["Min " + binning[0], "Max " + binning[0]]
                        )
                        ft = ft.drop(columns=["Min " + binning[0], "Max " + binning[0]])
                    except KeyError:
                        continue

            todump = todump.dropna(subset=["Max " + x_lab])
            ft = ft.dropna(subset=["Max " + x_lab])
            ex_outpath = os.path.join(
                self.excel_path, self.testname + "_" + x_ax + "_CE_tables.xlsx"
            )

            # Create a Pandas Excel writer using XlsxWriter as the engine.
            with pd.ExcelWriter(ex_outpath, engine="xlsxwriter") as writer:
                # dump global table
                todump = todump[
                    [
                        "Min " + x_lab,
                        "Max " + x_lab,
                        "C/E",
                        "Standard Deviation (σ)",
                    ]
                ]

                todump.to_excel(writer, sheet_name="Global")
                col_min = x_lab + "-min " + "[" + MCNP_UNITS[x_ax] + "]"
                col_max = x_lab + "-max " + "[" + MCNP_UNITS[x_ax] + "]"
                # Elaborate table for better output format

                ft[col_min] = ft["Min " + x_lab]
                ft[col_max] = ft["Max " + x_lab]

                ft["C/E (mean +/- σ)"] = (
                    ft["C/E"].round(2).astype(str)
                    + " +/- "
                    + ft["Standard Deviation (σ)"].round(2).astype(str)
                )
                # Delete all confusing columns
                for column in [
                    "Min " + x_lab,
                    "Max " + x_lab,
                    "C/E",
                    "Standard Deviation (σ)",
                ]:
                    del ft[column]

                # Dump also table material by material
                for input in self.inputs:
                    # dump material table
                    todump = ft.loc[input]

                    todump = todump.pivot(
                        index=["Quantity", col_min, col_max],
                        columns="Library",
                        values="C/E (mean +/- σ)",
                    )

                    todump.sort_values(by=[col_min])

                    todump.to_excel(writer, sheet_name=input, startrow=2)
                    ws = writer.sheets[input]
                    if skipcol_global == 0:
                        ws.write_string(0, 0, '"C/E (mean +/- σ)"')

                    # adjust columns' width
                    writer.sheets[input].set_column(0, 4, 18)

        return

    def _data_collect(self, input, tallynum, quantity_CE, e_intervals):
        """Collect data for C/E tables

        Parameters
        ----------
        input : str
            String containing the name of the benchmark case being run, will
            be subfolder test name for benchmarks with multiple runs
        tallynum : int
            tally number to be printed to table
        quantity_CE : str
            String containing the Y axis variable from benchmark config file
        e_intervals : list
            list of energy intervals from experimental benchmark config file

        Returns
        -------
        data : list
            list of dictionaries containing the data to be printed to the table
        x_lab : str
            Name of quantity being compared (not utilised?)

        """
        if self.multiplerun:
            filename = self.testname + "_" + input + "_" + str(tallynum)
        else:
            filename = self.testname + "_" + str(tallynum)
        # check if correspondent experimental data exists
        try:
            col_idx = self.exp_results[input][filename].columns.tolist()
        except KeyError:
            return None, None

        x_lab = col_idx[0]
        y_lab = col_idx[1]
        x = self.exp_results[input][filename][col_idx[0]].values
        y = self.exp_results[input][filename][col_idx[1]].values
        err = self.exp_results[input][filename][col_idx[2]].values
        # lib will be passed to the plotter
        lib = {"x": x, "y": y, "err": err, "ylabel": "Experiment"}
        # Get also the interpolator
        interpolator = interp1d(x, y, fill_value=0, bounds_error=False)
        # Collect the data to be plotted
        data = [lib]  # The first one should be the exp one
        for lib_tag in self.lib[1:]:  # Avoid exp
            lib_name = self.session.conf.get_lib_name(lib_tag)
            try:  # The tally may not be defined
                # Data for the plotter
                values = self.results[input, lib_tag][tallynum]
                lib = {
                    "x": values[x_lab],
                    "y": values["C"],
                    "err": values["Error"],
                    "ylabel": lib_name,
                }
                data.append(lib)
                # data for the table
                table = _get_tablevalues(
                    values, interpolator, x=x_lab, e_intervals=e_intervals
                )
                table["Quantity"] = quantity_CE
                table["Input"] = input
                table["Library"] = lib_name
                self.tables.append(table)
            except KeyError:
                # The tally is not defined
                pass
        return data, x_lab

    def _pp_excel_comparison(self):
        # Excel is actually printed by the build atlas in this case
        pass

    def _processMCNPdata(self, output):
        """
        given the mctal file the lethargy flux and energies are returned
        both for the neutron and photon tally

        Parameters
        ----------
        output : MCNPoutput
            object representing the MCNP output.

        Returns
        -------
        res : dic
            contains the extracted lethargy flux and energies.

        """
        res = {}
        # Read tally energy binned fluxes
        for tallynum, data in output.tallydata.items():
            tallynum = str(tallynum)
            res2 = res[tallynum] = {}
            x_axis = data.columns.tolist()[0]

            # Delete the total value
            data = data.set_index(x_axis).drop("total").reset_index()
            flux, energies, errors = self._parse_data_df(data, output, x_axis, tallynum)

            res2[x_axis + " [" + MCNP_UNITS[x_axis] + "]"] = energies
            res2["C"] = flux
            res2["Error"] = errors

            res[tallynum] = res2

        return res

    def _parse_data_df(self, data, output, x_axis, tallynum):
        """Read information from data DataFrame

        Parameters
        ----------
        data : DataFrame
            DataFrame containing all tally data for an output
        output : JADE MCNPoutput
            MCNP output object generated by MCNP parser
        x_axis : str
            X axis title
        tallynum : int
            Tally number, used to determine behaviour for protons and
            neutrons

        Returns
        -------
        flux : list
            list of binned flux values
        energies: list
            list of energy bin boundaries
        errors : list
            list of binned error values
        """
        # Generate a folder for each library
        flux = data["Value"].values
        energies = data[x_axis].values
        errors = data["Error"].values

        if TALLY_NORMALIZATION[self.testname] == "lethargy":
            # Energies for lethargy computation
            ergs = [1e-10]  # Additional "zero" energy for lethargy computation
            ergs.extend(energies.tolist())
            ergs = np.array(ergs)

            # Different behaviour for photons and neutrons
            for tally in output.mctal.tallies:
                if tallynum == str(tally.tallyNumber):
                    particle = tally.particleList[
                        np.where(tally.tallyParticles == 1)[0][0]
                    ]
            if particle == "Neutron":
                flux = flux / np.log((ergs[1:] / ergs[:-1]))
            elif particle == "Photon":
                flux = flux / (ergs[1:] - ergs[:-1])

        elif TALLY_NORMALIZATION[self.testname] == "energy bins":
            # Energies for lethargy computation
            data["bin"] = None

            prev_e = 0

            for e in data[x_axis].unique().tolist():
                data.loc[data[x_axis] == e, "bin"] = e - prev_e
                prev_e = e
            flux = flux / data["bin"].values
        return flux, energies, errors


def _get_tablevalues(
    df, interpolator, x="Energy [MeV]", y="C", e_intervals=[0.1, 1, 5, 10, 20]
):
    """
    Given the benchmark and experimental results returns a df to compile the
    C/E table for energy intervals

    Parameters
    ----------
    df : dict
        benchmark data.
    interpolator : func
        interpolator from experimental data.
    x : str, optional
        x column. The default is 'Energy [MeV]'.
    y : str, optional
        y columns. The default is 'C'.
    e_intervals : list, optional
        energy intervals to be used. The default is [0, 0.1, 1, 5, 10, 20].

    Returns
    -------
    pd.DataFrame
        C/E table per energy interval.

    """
    rows = []
    df = pd.DataFrame(df)
    df["Exp"] = interpolator(df[x])
    df["C/E"] = df[y] / df["Exp"]
    # it is better here to drop inf values because it means that the
    # interpolated experiment was zero, i.e., no value available
    df.replace([np.inf, -np.inf], np.nan, inplace=True)  # replace inf with NaN
    df.dropna(subset=["C/E"], how="all", inplace=True)  # drop the inf rows

    e_min = e_intervals[0]
    for e_max in e_intervals[1:]:
        red = df[e_min < df[x]]
        red = red[red[x] < e_max]
        mean = red["C/E"].mean()
        std = red["C/E"].std()
        row = {
            "C/E": mean,
            "Standard Deviation (σ)": std,
            "Max " + x[0]: e_max,
            "Min " + x[0]: e_min,
        }
        rows.append(row)
        # adjourn min energy
        e_min = e_max

    return pd.DataFrame(rows)


class TiaraOutput(ExperimentalOutput):
    def _processMCNPdata(self, output):
        return None

    def _case_tree_df_build(self):
        """
        Builds a dataframe containing library, source energy, shield material
        and thickness for each benchmark case, with all tallies for each case

        Returns
        -------
        pd.Dataframe
            DataFrame containing details about each benchmark case and the
            output tallies for that case
        """
        to_concat = []

        # Loop over libraries
        for lib in self.lib[1:]:
            # Declare empty dataframes
            case_tree = pd.DataFrame()
            for cont, case in enumerate(self.inputs):
                # Get data from benchmark's cases' names
                mat_name_list = case.split("-")
                if mat_name_list[0] == "cc":
                    case_tree.loc[cont, "Shield Material"] = "Concrete"
                elif mat_name_list[0] == "fe":
                    case_tree.loc[cont, "Shield Material"] = "Iron"
                case_tree.loc[cont, "Energy"] = int(mat_name_list[1])
                case_tree.loc[cont, "Shield Thickness"] = int(mat_name_list[2])
                case_tree.loc[cont, "Library"] = self.session.conf.get_lib_name(lib)
                # Put tally values in dataframe
                for tally in self.outputs[(case, lib)].mctal.tallies:
                    temp = (self.raw_data)[(case, lib)]
                    val = temp[tally.tallyNumber].iloc[-1]["Value"]
                    err = temp[tally.tallyNumber].iloc[-1]["Error"]
                    case_tree.loc[cont, tally.tallyComment] = val
                    case_tree.loc[cont, str(tally.tallyComment[0]) + " Error"] = err
            # Sort data in dataframe and assign to variable
            indexes = ["Library", "Shield Material", "Energy", "Shield Thickness"]
            case_tree.sort_values(indexes, inplace=True)
            case_tree = case_tree.set_index(indexes)
            case_tree.index.names = indexes
            # Add to overall dataframe
            to_concat.append(case_tree)
        # Return complete dataframe
        return pd.concat(to_concat)

    def _exp_comp_case_check(self, indexes):
        """
        Removes from mcnp case dataframe experimental data which don't have
        correspondent mcnp outputs and removes mcnp outputs without
        corresponding experimental data.

        Parameters
        ----------
        indexes : list
            Set of indexes to build the MultiIndex structure of the dataframes
        """
        self.case_tree_df = self.case_tree_df.reset_index()
        self.case_tree_df = self.case_tree_df.set_index(indexes[1:])
        # Delete experimental data
        com_index = self.case_tree_df.index.intersection(self.exp_data.index)
        self.exp_data = self.exp_data[self.exp_data.index.isin(com_index)]
        self.case_tree_df = self.case_tree_df[self.case_tree_df.index.isin(com_index)]
        self.case_tree_df = self.case_tree_df.reset_index()
        self.case_tree_df = self.case_tree_df.set_index(indexes)
        return

    def _get_conv_df(self, df):
        """
        Adds extra columns to the dataframe containing the maximum and
        average errors of the tallies

        Parameters
        ----------
        df : DataFrame
            DataFrame containing details about each benchmark case and the
            output tallies for that case

        Returns
        -------
        conv_df: Dataframe
            Same as previous dataframe, but with two extra columns containing
            maximum and average errors

        """
        conv_df = pd.DataFrame()
        for library in df.index.unique(level="Library").tolist():
            lib_df = df.loc[library]
            lib_err_df = lib_df.iloc[:, 1::2]
            max = lib_err_df.max().max()
            avg = lib_err_df.mean().mean()
            conv_df.loc["Max Error", library] = max
            conv_df.loc["Average Error", library] = avg
        return conv_df


class TiaraFCOutput(TiaraOutput):
    def _pp_excel_comparison(self):
        """
        Builds dataframe from computational output comparable to experimental
        data and generates the excel comparison
        """

        # Get computational data structure for each library
        self.case_tree_df = self._case_tree_df_build()

        off_dict = {0: "On-axis", 20: "20 cm off-axis"}
        columns = ["U238", "U238 Error", "Th232", "Th232 Error"]
        new_idx_list = []
        # build computational dataframe with the same index structure as exp
        # dataframe
        for idx in self.case_tree_df.index.values.tolist():
            for offset in [0, 20]:
                new_idx = idx + (offset,)
                new_idx_list.append(new_idx)
        indexes = [
            "Library",
            "Shield Material",
            "Energy",
            "Shield Thickness",
            "Axis offset",
        ]
        multi_index = pd.MultiIndex.from_tuples(new_idx_list, names=indexes)
        case_tree_df_2 = pd.DataFrame(index=multi_index, columns=columns)
        # Sort to avoid later warnings
        case_tree_df_2.sort_values(indexes, axis=0, inplace=True)
        self.case_tree_df.sort_values(indexes[:-1], axis=0, inplace=True)
        # Put values into new dataframe
        for idx in self.case_tree_df.index.values.tolist():
            for offset in [0, 20]:
                for err_string in ["", " Error"]:
                    val_str = off_dict[offset] + " 238U FC" + err_string
                    val = self.case_tree_df.loc[idx, val_str]
                    case_tree_df_2.loc[idx + (offset,), "U238" + err_string] = val
                    val_str = off_dict[offset] + " 232Th FC" + err_string
                    val = self.case_tree_df.loc[idx, val_str]
                    case_tree_df_2.loc[idx + (offset,), "Th232" + err_string] = val

        self.case_tree_df = case_tree_df_2.copy()
        # Discard experimental data without a correspondent computational data
        self._exp_comp_case_check(indexes=indexes)
        self.case_tree_df.sort_values(indexes, axis=0, inplace=True)
        # Build ExcelWriter object
        filepath = os.path.join(self.excel_path, "Tiara_Fission_Cells_CE_tables.xlsx")
        with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
            # Create 1 worksheet for each energy/material combination
            mats = self.case_tree_df.index.unique(level="Shield Material").tolist()
            ens = self.case_tree_df.index.unique(level="Energy").tolist()
            for shield_material in mats:
                for energy in ens:
                    # Set MultiIndex structure of the table
                    # Set column names
                    column_names = []
                    temp_df = self.case_tree_df.loc(axis=0)[
                        :, shield_material, energy
                    ].copy()
                    for fission_cell in ["U238", "Th232"]:
                        column_names.append(("Exp", fission_cell, "Value"))
                        column_names.append(("Exp", fission_cell, "Error"))
                    libs = self.case_tree_df.index.unique(level="Library").tolist()
                    for lib in libs:
                        for fission_cell in ["U238", "Th232"]:
                            column_names.append((lib, fission_cell, "Value"))
                            column_names.append((lib, fission_cell, "C/E"))
                            column_names.append((lib, fission_cell, "C/E Error"))
                    names = ["Library", "Fission Cell", ""]
                    column_index = pd.MultiIndex.from_tuples(column_names, names=names)
                    # Set row indexes
                    row_idx_list = []
                    for idx in temp_df.index.values.tolist():
                        row_idx_list.append((idx[-2], idx[-1]))
                    names = ["Shield Thickness", "Axis offset"]
                    row_idx = pd.MultiIndex.from_tuples(row_idx_list, names=names)

                    # Build new dataframe with desired multindex structure
                    new_dataframe = pd.DataFrame(columns=column_index, index=row_idx)
                    # Fill the new dataframe with proper values
                    for idx_row in new_dataframe.index.values.tolist():
                        for idx_col in new_dataframe.columns.values.tolist():
                            row_tuple = (
                                shield_material,
                                energy,
                                idx_row[0],
                                idx_row[1],
                            )
                            if idx_col[0] == "Exp":
                                if idx_col[2] == "Value":
                                    val = self.exp_data.loc[row_tuple, idx_col[1]]
                                    new_dataframe.loc[idx_row, idx_col] = val
                                else:
                                    val = self.exp_data.loc[
                                        row_tuple, idx_col[1] + " Error"
                                    ]
                                    new_dataframe.loc[idx_row, idx_col] = val
                            else:
                                row_tuple = (idx_col[0],) + row_tuple
                                if idx_col[2] == "Value":
                                    val = temp_df.loc[row_tuple, idx_col[1]]
                                    new_dataframe.loc[idx_row, idx_col] = val
                                elif idx_col[2] == "C/E Error":
                                    val1 = temp_df.loc[row_tuple, idx_col[1] + " Error"]
                                    val2 = self.exp_data.loc[
                                        row_tuple[1:], idx_col[1] + " Error"
                                    ]
                                    ce_err = math.sqrt(val1**2 + val2**2)
                                    new_dataframe.loc[idx_row, idx_col] = ce_err
                                else:
                                    val = temp_df.loc[row_tuple, idx_col[1]]
                                    val2 = self.exp_data.loc[row_tuple[1:], idx_col[1]]
                                    new_dataframe.loc[idx_row, idx_col] = val / val2
                    # Assign worksheet title and put into Excel
                    conv_df = self._get_conv_df(temp_df)
                    sheet_name = "Tiara FC {}, {} MeV".format(
                        shield_material, str(energy)
                    )
                    sort = ["Axis offset", "Shield Thickness"]
                    new_dataframe.sort_values(sort, axis=0, inplace=True)
                    new_dataframe = new_dataframe.drop_duplicates()
                    new_dataframe.to_excel(writer, sheet_name=sheet_name)
                    conv_df.to_excel(writer, sheet_name=sheet_name, startrow=18)

    def _read_exp_results(self):
        """
        Reads and manipulates conderc Excel file
        """

        # Read experimental data from CONDERC Excel file
        filepath = os.path.join(
            self.path_exp_res, "FC_BS_Experimental-results-CONDERC.xlsx"
        )
        FC_data = {
            ("Iron", "43"): pd.read_excel(
                filepath, sheet_name="Fission cell", usecols="A:E", skiprows=2, nrows=10
            ),
            ("Iron", "68"): pd.read_excel(
                filepath,
                sheet_name="Fission cell",
                usecols="A:E",
                skiprows=16,
                nrows=10,
            ),
            ("Concrete", "43"): pd.read_excel(
                filepath, sheet_name="Fission cell", usecols="A:E", skiprows=30, nrows=8
            ),
            ("Concrete", "68"): pd.read_excel(
                filepath, sheet_name="Fission cell", usecols="A:E", skiprows=42, nrows=8
            ),
        }
        # Build experimental dataframe
        index = ["Shield Material", "Energy", "Shield Thickness", "Axis offset"]
        to_concat = []
        for idx, element in FC_data.items():
            # Build a first useful structure from CONDERC data
            element["Shield Material"] = idx[0]
            element["Energy"] = int(idx[1])
            element[["Shield Thickness", "Axis offset"]] = element[
                "Fission c./Position (shield t., axis offset)"
            ].str.split(",", expand=True)
            element["Shield Thickness"] = element["Shield Thickness"].astype("int")
            element["Axis offset"] = element["Axis offset"].astype("int")
            element.drop(
                "Fission c./Position (shield t., axis offset)", axis=1, inplace=True
            )

            element = element.set_index(index)
            element.index.names = index
            to_concat.append(element)
        exp_data = pd.concat(to_concat)
        # Make exp data normalization compatible with tally outputs
        exp_data["238 U [/1e24]"] *= 1e24
        exp_data["232 Th [/1e24]"] *= 1e24
        exp_data["err [%]"] /= 100
        exp_data["err [%].1"] /= 100
        # Rename columns and sort values
        exp_data.columns = ["U238", "U238 Error", "Th232", "Th232 Error"]
        inde = ["Shield Material", "Energy", "Axis offset", "Shield Thickness"]
        exp_data.sort_values(inde, axis=0, inplace=True)
        # Assign exp data variable
        self.exp_data = exp_data

    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str
            path to the temporary folder containing the plots for the atlas
        atlas : Atlas
            Object representing the plot Atlas.
        """
        # Set plot and axes details
        unit = "-"
        quantity = ["On-axis reaction rate", "Off-axis 20 cm reaction rate"]
        xlabel = "Shield thickness [cm]"
        f_cell_list = ["U238", "Th232"]
        # Loop over shield material/energy combinations
        mat_list = self.case_tree_df.index.unique(level="Shield Material").tolist()
        e_list = self.case_tree_df.index.unique(level="Energy").tolist()

        for shield_material in mat_list:
            for energy in e_list:
                # Put proper data in data dict to be sent to plotter
                data_U_p = []
                data_Th_p = []
                exp_dat = self.exp_data.loc(axis=0)[shield_material, energy]
                thick_list = exp_dat.index.unique(level="Shield Thickness").tolist()
                off_list = exp_dat.index.unique(level="Axis offset").tolist()
                for thick in thick_list:
                    for offset in off_list:
                        idx = (thick, offset)
                        if idx not in exp_dat.index.values.tolist():
                            exp_dat.loc[idx, :] = [None] * len(exp_dat.columns)
                exp_dat.replace(to_replace=[None], value=np.nan, inplace=True)
                sort_idx = ["Axis offset", "Shield Thickness"]
                exp_dat.sort_values(sort_idx, axis=0, inplace=True)
                x = np.array(thick_list)

                y = {}
                err = {}
                # Put data in proper lists
                for f_cell in f_cell_list:
                    y[f_cell] = []
                    err[f_cell] = []

                for f_cell in f_cell_list:
                    for offset in off_list:
                        y_dat = exp_dat.loc(axis=0)[:, offset]
                        y[f_cell].append(y_dat[f_cell].to_numpy())
                        err[f_cell].append(y_dat[f_cell + " Error"].to_numpy())
                # Append experimental data
                ylabel = "Experiment"
                data_U = {"x": x, "y": y["U238"], "err": err["U238"], "ylabel": ylabel}
                data_Th = {
                    "x": x,
                    "y": y["Th232"],
                    "err": err["Th232"],
                    "ylabel": ylabel,
                }
                data_U_p.append(data_U)
                data_Th_p.append(data_Th)

                # Loop over libraries
                for lib in self.lib[1:]:
                    # Get proper computational data
                    ylabel = self.session.conf.get_lib_name(lib)
                    mcnp_data = self.case_tree_df.loc(axis=0)[
                        ylabel, shield_material, energy
                    ]
                    thick_list = mcnp_data.index.unique(
                        level="Shield Thickness"
                    ).tolist()
                    off_list = mcnp_data.index.unique(level="Axis offset").tolist()
                    for thick in thick_list:
                        for offset in off_list:
                            idx = (thick, offset)
                            if idx not in mcnp_data.index.values.tolist():
                                col = mcnp_data.columns
                                mcnp_data.loc[idx, :] = [None] * len(col)
                    mcnp_data.replace(to_replace=[None], value=np.nan, inplace=True)
                    mcnp_data.sort_values(
                        ["Axis offset", "Shield Thickness"], axis=0, inplace=True
                    )
                    # Save proper computational data in variables
                    x = np.array(thick_list)

                    y = {}
                    err = {}
                    # Put data in proper lists
                    for f_cell in f_cell_list:
                        y[f_cell] = []
                        err[f_cell] = []
                    for f_cell in f_cell_list:
                        for offset in off_list:
                            y_dat = mcnp_data.loc(axis=0)[:, offset]
                            y[f_cell].append(y_dat[f_cell].to_numpy())
                            err[f_cell].append(y_dat[f_cell + " Error"].to_numpy())

                    # Append computational data to data dict
                    data_U = {
                        "x": x,
                        "y": y["U238"],
                        "err": err["U238"],
                        "ylabel": ylabel,
                    }
                    data_Th = {
                        "x": x,
                        "y": y["Th232"],
                        "err": err["Th232"],
                        "ylabel": ylabel,
                    }
                    data_U_p.append(data_U)
                    data_Th_p.append(data_Th)
                    fission_cell = ["Uranium-238", "Thorium-232"]

                for cont, data in enumerate([data_U_p, data_Th_p]):
                    # Set title and send to plotter
                    title = "Tiara Experiment: {} Fission Cell detector, Energy: {} MeV, Shield material: {}".format(
                        fission_cell[cont], str(energy), shield_material
                    )
                    hea = atlas.doc.add_heading(title, level=1)
                    hea.alignment = 1
                    outname = "tmp"
                    plot = Plotter(
                        data,
                        title,
                        tmp_path,
                        outname,
                        quantity,
                        unit,
                        xlabel,
                        self.testname,
                    )
                    img_path = plot.plot("Waves")
                    atlas.insert_img(img_path, width=Inches(9))
                    atlas.doc.add_page_break()
        return atlas


class TiaraBSOutput(TiaraOutput):
    def _pp_excel_comparison(self):
        """
        This method prints Tiara C/E tables for Bonner Spheres detectors
        """

        # Get main dataframe with computational data of all cases
        self.case_tree_df = self._case_tree_df_build()
        # Rename columns of mcnp dataframe properly
        columns = ["Bare", "15 mm", "30 mm", "50 mm", "90 mm"]
        err_columns = []
        for strings in columns:
            err_columns.append(strings + " Error")
        columns_mcnp = []
        for i, col in enumerate(columns):
            columns_mcnp.append(columns[i])
            columns_mcnp.append(err_columns[i])

        self.case_tree_df.rename(
            columns=dict(zip(self.case_tree_df.columns, columns_mcnp)), inplace=True
        )
        indexes = ["Library", "Shield Material", "Energy", "Shield Thickness"]
        self._exp_comp_case_check(indexes=indexes)
        # Create ExcelWriter object
        filepath = os.path.join(self.excel_path, "Tiara_Bonner_Spheres_CE_tables.xlsx")
        with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
            # Loop over shield material/energy combinations
            mat_list = self.case_tree_df.index.unique(level="Shield Material").tolist()
            e_list = self.case_tree_df.index.unique(level="Energy").tolist()
            for shield_material in mat_list:
                for energy in e_list:
                    # Select the cases with the energy/material combination
                    column_names = []
                    comp_data = self.case_tree_df.loc(axis=0)[
                        :, shield_material, energy
                    ]
                    exp_data = self.exp_data.loc(axis=0)[shield_material, energy]
                    thick_list = exp_data.index.unique().tolist()
                    for shield_thickness in thick_list:
                        column_names.append(("Exp", shield_thickness, "Value"))
                    lib_list = comp_data.index.unique(level="Library").tolist()
                    for lib in lib_list:
                        thick_list = comp_data.index.unique(
                            level="Shield Thickness"
                        ).tolist()
                        for shield_thickness in thick_list:
                            column_names.append((lib, shield_thickness, "Value"))
                            column_names.append((lib, shield_thickness, "Error"))
                            column_names.append((lib, shield_thickness, "C/E"))
                    names = ["Library", "Shield Thickness", ""]
                    index = pd.MultiIndex.from_tuples(column_names, names=names)

                    # Create new dataframe with the MultiIndex structure
                    new_dataframe = pd.DataFrame(index=columns, columns=index)

                    # Add the proper values in the new dataframe
                    for idx_row in new_dataframe.index.values.tolist():
                        for idx_col in new_dataframe.columns.values.tolist():
                            if idx_col[0] == "Exp":
                                val = exp_data.loc[idx_col[1], idx_row]
                                new_dataframe.loc[idx_row, idx_col] = val
                            else:
                                row_tuple = (
                                    idx_col[0],
                                    shield_material,
                                    energy,
                                    idx_col[1],
                                )
                                if idx_col[2] == "Value":
                                    val = comp_data.loc[row_tuple, idx_row]
                                    new_dataframe.loc[idx_row, idx_col] = val
                                elif idx_col[2] == "Error":
                                    val = comp_data.loc[row_tuple, idx_row + " Error"]
                                    new_dataframe.loc[idx_row, idx_col] = val
                                else:
                                    val = comp_data.loc[row_tuple, idx_row]
                                    val2 = exp_data.loc[idx_col[1], idx_row]
                                    new_dataframe.loc[idx_row, idx_col] = val / val2

                    # Print the dataframe in a worksheet in Excel file
                    conv_df = self._get_conv_df(comp_data)
                    sheet_name = "Tiara {}, {} MeV".format(shield_material, str(energy))
                    new_dataframe.to_excel(writer, sheet_name=sheet_name)
                    conv_df.to_excel(writer, sheet_name=sheet_name, startrow=12)

    def _read_exp_results(self):
        """
        Reads and manipulates conderc Excel file

        """
        # Get experimental data filepath
        filepath = os.path.join(
            self.path_exp_res, "FC_BS_Experimental-results-CONDERC.xlsx"
        )
        # Read exp data from CONDERC excel file
        s_name = "Bonner sphere"
        BS_data = {
            ("Iron", "43"): pd.read_excel(
                filepath, sheet_name=s_name, usecols="A:F", skiprows=2, nrows=3
            ),
            ("Iron", "68"): pd.read_excel(
                filepath, sheet_name=s_name, usecols="A:F", skiprows=9, nrows=3
            ),
            ("Concrete", "43"): pd.read_excel(
                filepath, sheet_name=s_name, usecols="A:F", skiprows=16, nrows=4
            ),
            ("Concrete", "68"): pd.read_excel(
                filepath, sheet_name=s_name, usecols="A:F", skiprows=24, nrows=3
            ),
        }

        for key, value in BS_data.items():
            value["Shield Material"] = key[0]
            value["Energy"] = int(key[1])

        to_concat = []
        for value in BS_data.values():
            to_concat.append(value)
        exp_data = pd.concat(to_concat, ignore_index=True)

        # Adjust experimental data dataframe's structure
        exp_data.rename(
            columns={"Polyethylene t./Shield t.": "Shield Thickness"}, inplace=True
        )
        indexes = ["Shield Material", "Energy", "Shield Thickness"]
        exp_data = exp_data.set_index(indexes)
        exp_data.sort_values(indexes, axis=0, inplace=True)

        # Save experimental data
        self.exp_data = exp_data

    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str
            path to the temporary folder containing the plots for the atlas
        atlas : Atlas
            Object representing the plot Atlas.
        """
        # Set plot axes
        unit = "-"
        quantity = ["Experiment reaction rate"]
        xlabel = "Bonner Sphere Radius [mm]"
        x = ["Bare", "15", "30", "50", "90"]

        # Loop over all benchmark cases (materials)
        for idx, values in self.exp_data.iterrows():
            data = []
            # Get experimental data and errors for the selected benchmark case
            y = [values]
            err = [np.zeros(len(y))]

            # Append experimental data to data list (sent to plotter)
            ylabel = "Experiment"
            data_p = {"x": x, "y": y, "err": err, "ylabel": ylabel}
            data.append(data_p)

            # Loop over selected libraries
            for lib in self.lib[1:]:
                # Get library name, assign title to the plot
                ylabel = self.session.conf.get_lib_name(lib)
                title = (
                    "Tiara Experiment: Bonner Spheres detector, Energy: "
                    + str(idx[1])
                    + " MeV, Shield material: "
                    + idx[0]
                    + ", Shield thickness: "
                    + str(idx[2])
                    + " cm"
                )

                # Get computational data and errors for the selected case
                comp_idx = (ylabel,) + idx
                y = [self.case_tree_df.loc[comp_idx, ::2]]
                err = self.case_tree_df.loc[comp_idx]
                err = [err.iloc[1::2]]

                # Append computational data to data list(to be sent to plotter)
                data_p = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                data.append(data_p)

            # Send data to plotter
            hea = atlas.doc.add_heading(title, level=1)
            hea.alignment = 1
            outname = "tmp"
            plot = Plotter(
                data, title, tmp_path, outname, quantity, unit, xlabel, self.testname
            )
            img_path = plot.plot("Waves")
            atlas.insert_img(img_path, width=Inches(9))
            atlas.doc.add_page_break()

        return atlas


class ShieldingOutput(ExperimentalOutput):
    def _processMCNPdata(self, output):
        return None

    def _pp_excel_comparison(self):
        """
        This method prints C/E tables for shielding benchmark comparisons
        """
        # FNG SiC specific corrections/normalisations
        fngsic_k = [0.212, 0.204, 0.202, 0.202]  # Neutron sensitivity of TL detectors
        fngsic_norm = 1.602e-13 * 1000  # J/MeV * g/kg
        lib_names_dict = {}
        column_names = []
        column_names.append(("Exp", "Value"))
        column_names.append(("Exp", "Error"))
        for lib in self.lib[1:]:
            namelib = self.session.conf.get_lib_name(lib)
            lib_names_dict[namelib] = lib
            column_names.append((namelib, "Value"))
            column_names.append((namelib, "C/E"))
            column_names.append((namelib, "C/E Error"))

        names = ["Library", ""]
        column_index = pd.MultiIndex.from_tuples(column_names, names=names)
        # filepath = self.excel_path_mcnp + '\\' + self.testname + '_CE_tables.xlsx'
        filepath = os.path.join(self.excel_path, f"{self.testname}_CE_tables.xlsx")
        with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
            # TODO Replace when other transport codes implemented.
            code = "mcnp"
            for mat in self.inputs:
                exp_folder = os.path.join(self.path_exp_res, mat)
                exp_filename = self.testname + "_" + mat + ".csv"
                exp_filepath = os.path.join(exp_folder, exp_filename)
                exp_data_df = pd.read_csv(exp_filepath)
                # Get experimental data and errors for the selected benchmark case
                x = exp_data_df["Depth"].values.tolist()
                indexes = pd.Index(data=x, name="Depth [cm]")
                df_tab = pd.DataFrame(index=indexes, columns=column_index)
                for idx_col in df_tab.columns.values.tolist():
                    if idx_col[0] == "Exp":
                        if idx_col[1] == "Value":
                            vals = exp_data_df.loc[:, "Reaction Rate"].tolist()
                            df_tab[idx_col] = vals
                        else:
                            vals = exp_data_df.loc[:, "Error"].to_numpy() / 100
                            vals = vals.tolist()
                            df_tab[idx_col] = vals
                    else:
                        t = (mat, lib_names_dict[idx_col[0]])
                        if idx_col[1] == "Value":
                            if mat != "TLD":
                                vals = self.raw_data[t][4]["Value"].values[: len(x)]
                            else:
                                # FNG SiC experiment measured the total dose
                                if self.testname == "FNG-SiC":
                                    # Neutron dose
                                    Dn = (
                                        self.raw_data[t][16]["Value"].values[
                                            : len(x)
                                        ]
                                    ) * fngsic_norm
                                    Dn_multiplied = [
                                        value * constant
                                        for value, constant in zip(Dn, fngsic_k)
                                    ]
                                    # Photon dose
                                    Dp = (
                                        self.raw_data[t][26]["Value"].values[
                                            : len(x)
                                        ]
                                    ) * fngsic_norm
                                    # Sum neutron and photon dose with neutron sensitivity as a function of depth
                                    Dt = [sum(pair) for pair in zip(Dn_multiplied, Dp)]
                                    vals = Dt
                                else:
                                    vals = self.raw_data[t][6]["Value"].values[
                                        : len(x)
                                    ]
                            df_tab[idx_col] = vals
                        elif idx_col[1] == "C/E Error":
                            if mat != "TLD":
                                errs = self.raw_data[t][4]["Error"].values[: len(x)]
                            else:
                                if self.testname == "FNG-SiC":
                                    errs = np.sqrt(
                                        np.square(
                                            self.raw_data[t][16]["Error"].values[
                                                : len(x)
                                            ]
                                        )
                                        + np.square(
                                            self.raw_data[t][26]["Error"].values[
                                                : len(x)
                                            ]
                                        )
                                    )
                                else:
                                    errs = self.raw_data[t][6]["Error"].values[
                                        : len(x)
                                    ]
                            vals1 = np.square(errs)
                            vals2 = np.square(
                                exp_data_df.loc[:, "Error"].to_numpy() / 100
                            )
                            ce_err = np.sqrt(vals1 + vals2)
                            ce_err = ce_err.tolist()
                            df_tab[idx_col] = ce_err
                        else:
                            if mat != "TLD":
                                vals1 = self.raw_data[t][4]["Value"].values[: len(x)]
                            else:
                                if self.testname == "FNG-SiC":
                                    # Neutron dose
                                    Dn = (
                                        self.raw_data[t][16]["Value"].values[
                                            : len(x)
                                        ]
                                    ) * fngsic_norm
                                    Dn_multiplied = [
                                        value * constant
                                        for value, constant in zip(Dn, fngsic_k)
                                    ]
                                    # Photon dose
                                    Dp = (
                                        self.raw_data[t][26]["Value"].values[
                                            : len(x)
                                        ]
                                    ) * fngsic_norm
                                    # Sum neutron and photon dose with neutron sensitivity as a function of depth
                                    Dt = [sum(pair) for pair in zip(Dn_multiplied, Dp)]
                                    vals1 = Dt
                                else:
                                    vals1 = self.raw_data[t][6]["Value"].values[
                                        : len(x)
                                    ]
                            vals2 = exp_data_df.loc[:, "Reaction Rate"].to_numpy()
                            ratio = vals1 / vals2
                            ratio = ratio.tolist()
                            df_tab[idx_col] = vals1 / vals2

                # Assign worksheet title and put into Excel
                conv_df = self._get_conv_df(mat, len(x))
                sheet = self.testname.replace("-", " ")
                sheet_name = sheet + ", Foil {}".format(mat)
                df_tab.to_excel(writer, sheet_name=sheet_name)
                conv_df.to_excel(writer, sheet_name=sheet_name, startrow=18)

    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str
            path to the temporary folder containing the plots for the atlas
        atlas : Atlas
            Object representing the plot Atlas.
        """
        # FNG SiC specific corrections/normalisations
        fngsic_k = [0.212, 0.204, 0.202, 0.202]  # Neutron sensitivity of TL detectors
        fngsic_norm = 1.602e-13 * 1000  # J/MeV * g/kg
        # Set plot and axes details
        unit = "-"
        xlabel = "Shielding thickness [cm]"
        data = []
        # TODO Replace when other transport codes implemented.
        code = "mcnp"
        for material in tqdm(self.inputs, desc="Foil: "):
            data = []
            exp_folder = os.path.join(self.path_exp_res, material)
            exp_filename = self.testname + "_" + material + ".csv"
            exp_filepath = os.path.join(exp_folder, exp_filename)
            exp_data_df = pd.read_csv(exp_filepath)
            # Get experimental data and errors for the selected benchmark case
            x = exp_data_df["Depth"].values
            y = []
            err = []
            y.append(exp_data_df["Reaction Rate"].values)
            err.append(exp_data_df["Error"].values / 100)
            # Append experimental data to data list (sent to plotter)
            ylabel = "Experiment"
            data_exp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
            data.append(data_exp)

            if material != "TLD":
                title = self.testname + " experiment, Foil: " + FOILS_REACTION[material]
            else:
                if self.testname == "FNG-SiC":
                    title = (
                        self.testname
                        + " experiment, Total absorbed dose in TLD detectors"
                    )
                else:
                    title = (
                        self.testname
                        + " experiment, Gamma absorbed dose in TLD-300 detectors"
                    )
            # Loop over selected libraries
            # Loop over selected libraries
            for lib in self.lib[1:]:
                # Get library name, assign title to the plot
                ylabel = self.session.conf.get_lib_name(lib)
                y = []
                err = []
                if material != "TLD":
                    v = self.raw_data[(material, lib)][4]["Value"].values[: len(x)]
                else:
                    if self.testname == "FNG-SiC":
                        # Neutron dose
                        Dn = (
                            self.raw_data[(material, lib)][16]["Value"].values[
                                : len(x)
                            ]
                        ) * fngsic_norm
                        Dn_multiplied = [
                            value * constant for value, constant in zip(Dn, fngsic_k)
                        ]
                        # Photon dose
                        Dp = (
                            self.raw_data[(material, lib)][26]["Value"].values[
                                : len(x)
                            ]
                        ) * fngsic_norm
                        # Sum neutron and photon dose with neutron sensitivity as a function of depth
                        v = [sum(pair) for pair in zip(Dn_multiplied, Dp)]
                    else:
                        v = self.raw_data[(material, lib)][6]["Value"].values[
                            : len(x)
                        ]
                y.append(v)
                if material != "TLD":
                    v = self.raw_data[(material, lib)][4]["Error"].values[: len(x)]
                else:
                    if self.testname == "FNG-SiC":
                        v = np.sqrt(
                            np.square(
                                self.raw_data[(material, lib)][16][
                                    "Error"
                                ].values[: len(x)]
                            )
                            + np.square(
                                self.raw_data[(material, lib)][26][
                                    "Error"
                                ].values[: len(x)]
                            )
                        )
                    else:
                        v = self.raw_data[(material, lib)][6]["Error"].values[
                            : len(x)
                        ]
                err.append(v)
                # Append computational data to data list(to be sent to plotter)
                data_comp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                data.append(data_comp)

            # Send data to plotter
            outname = "tmp"
            if material != "TLD":
                quantity = [ACTIVATION_REACTION[material] + " Reaction Rate"]
            else:
                quantity = ["Absorbed dose"]
            atlas.doc.add_heading(title, level=1)
            plot = Plotter(
                data, title, tmp_path, outname, quantity, unit, xlabel, self.testname
            )
            img_path = plot.plot("Waves")
            atlas.insert_img(img_path, width=Inches(9))
            atlas.doc.add_page_break()

        return atlas

    def _get_conv_df(self, mat, size):
        # TODO Replace when other transport codes implemented.
        code = "mcnp"
        conv_df = pd.DataFrame()
        for lib in self.lib[1:]:
            if mat != "TLD":
                max = self.raw_data[(mat, lib)][4]["Error"].values[:size].max()
                avg = self.raw_data[(mat, lib)][4]["Error"].values[:size].mean()
            else:
                if self.testname == "FNG-SiC":
                    v = np.sqrt(
                        np.square(
                            self.raw_data[(mat, lib)][16]["Error"].values[:size]
                        )
                        + np.square(
                            self.raw_data[(mat, lib)][26]["Error"].values[:size]
                        )
                    )
                    max = np.max(v)
                    avg = np.mean(v)
                else:
                    max = (
                        self.raw_data[(mat, lib)][6]["Error"].values[:size].max()
                    )
                    avg = (
                        self.raw_data[(mat, lib)][6]["Error"].values[:size].mean()
                    )
            library = self.session.conf.get_lib_name(lib)
            conv_df.loc["Max Error", library] = max
            conv_df.loc["Average Error", library] = avg
        return conv_df


class MultipleSpectrumOutput(SpectrumOutput):
    def _build_atlas(self, tmp_path, atlas):
        """
        Fill the atlas with the customized plots. Creation and saving of the
        atlas are handled elsewhere.

        Parameters
        ----------
        tmp_path : str
            path to the temporary folder containing the plots for the atlas
        atlas : Atlas
            Object representing the plot Atlas.

        """
        self.tables = []
        self.groups = pd.read_excel(self.cnf_path)
        self.groups = self.groups.set_index(["Group", "Tally", "Input"])
        self.group_list = self.groups.index.get_level_values("Group").unique().tolist()
        for group in self.group_list:
            self._plot_tally_group(group, tmp_path, atlas)

        # Dump C/E table
        self._dump_ce_table()

        return atlas

    def _plot_tally_group(self, group, tmp_path, atlas):
        """
        Plots tallies for a given group of outputs and add to Atlas object

        Parameters
        ----------
        group : list
            list of groups in the experimental benchmark object, outputs are
            grouped by material, several tallies for each material/group
        tmp_path : str
            path to temporary atlas plot folder
        atlas : JADE Atlas
            Atlas object

        Returns
        -------
        atlas : JADE Atlas
            adjusted Atlas object
        """
        # Extract 'Tally' and 'Input' values for the current 'Group'
        group_data = self.groups.xs(group, level="Group", drop_level=False)
        data_group = {}
        group_lab = []
        mult_factors = group_data["Multiplying factor"].values.tolist()
        for m, idx in enumerate(group_data.index.tolist()):
            tallynum = idx[1]
            input = idx[2]
            if str(tallynum) not in self.results[input, self.lib[1]].keys():
                continue
            quantity = group_data.loc[(group, tallynum, input), "Quantity"]
            particle = group_data.loc[(group, tallynum, input), "Particle"]
            add_info = group_data.loc[(group, tallynum, input), "Y Label"]
            quant_string = particle + " " + quantity + " " + add_info
            e_int = group_data.loc[(group, tallynum, input), "C/E X Quantity intervals"]
            e_int = e_int.split("-")

            # Convert the list of number strings into a list of integers
            e_intervals = [float(num) for num in e_int]
            data_temp, xlabel = self._data_collect(
                input, str(tallynum), quant_string, e_intervals
            )
            if data_temp is None:
                continue
            data_group[m] = data_temp
            unit = group_data.loc[(group, tallynum, input), "Y Unit"]

            group_lab.append(add_info)
            # Once the data is collected it is passed to the plotter
        title = self._define_title(input, particle, quantity)
        outname = "tmp"
        plot = Plotter(
            data_group,
            title,
            tmp_path,
            outname,
            quantity,
            unit,
            xlabel,
            self.testname,
            group_num=group,
            add_labels=group_lab,
            mult_factors=mult_factors,
        )
        img_path = plot.plot("Experimental points group")
        atlas.doc.add_heading(title, level=1)
        atlas.insert_img(img_path)
        img_path = plot.plot("Experimental points group CE")
        atlas.doc.add_heading(title + " C/E", level=1)
        atlas.insert_img(img_path, width=Inches(9))
        return atlas

    def _define_title(self, input, particle, quantity):
        """
        determines which benchmark is being compared and assigns title
        accordinly

        Parameters
        ----------
        input : str
            Test name
        particle : str
            Particle being tallied
        quantity : str
            Type of quantity being plotted on the X axis

        Returns
        -------
        Title: str
            Title string
        """

        if not self.multiplerun:
            title = self.testname + ", " + particle + " " + quantity
        elif self.testname == "Tiara-BC":
            mat = input.split("-")[0]
            if mat == "cc":
                material = "Concrete"
            else:
                material = "Iron"
            energy = input.split("-")[1]
            sh_th = input.split("-")[2]
            add_coll = input.split("-")[3]
            title = (
                self.testname
                + ", Shielding: "
                + material
                + ", "
                + sh_th
                + "cm; Source energy: "
                + energy
                + " MeV; Additional collimator: "
                + add_coll
                + " cm"
            )
        elif self.testname == "FNS-TOF":
            mat = input.split("-")[0]
            sl_th = input.split("-")[1]
            title = self.testname + ", " + sl_th + "cm " + mat + " slab"
        else:
            title = self.testname + ", " + particle + " " + quantity
        return title


class fnghcpboutput(ExperimentalOutput):

    def _processMCNPdata(self, output):

        return None

    def _pp_excel_comparison(self):
        """Produces the Excel document for comparison to experiment."""

        lib_names_dict = {}
        column_names = []
        column_names.append(("Exp", "Value"))
        column_names.append(("Exp", "Error"))
        for lib in self.lib[1:]:
            namelib = self.session.conf.get_lib_name(lib)
            lib_names_dict[namelib] = lib
            column_names.append((namelib, "Value"))
            column_names.append((namelib, "C/E"))
            column_names.append((namelib, "C/E Error"))

        names = ["Library", ""]
        column_index = pd.MultiIndex.from_tuples(column_names, names=names)
        filepath = self.excel_path + "\\" + self.testname + "_CE_tables.xlsx"
        with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
            code = "mcnp"
            for mat in self.inputs:
                exp_folder = os.path.join(self.path_exp_res, mat)
                exp_filename = self.testname + "_" + mat + ".csv"
                exp_filepath = os.path.join(exp_folder, exp_filename)
                exp_data_df = pd.read_csv(exp_filepath)

                # Get experimental data and errors for the selected benchmark case
                if mat == "H3":
                    x = exp_data_df["Pellet"].values.tolist()
                    indexes = pd.Index(data=x, name="Pellet #")
                else:
                    x = exp_data_df["Depth"].values.tolist()
                    indexes = pd.Index(data=x, name="Depth [cm]")

                df_tab = pd.DataFrame(index=indexes, columns=column_index)
                for idx_col in df_tab.columns.values.tolist():
                    if idx_col[0] == "Exp":
                        if idx_col[1] == "Value":
                            if mat == "H3":
                                vals = exp_data_df.loc[:, "Activity"].tolist()
                            else:
                                vals = exp_data_df.loc[:, "Reaction Rate"].tolist()
                            df_tab[idx_col] = vals
                        else:
                            vals = exp_data_df.loc[:, "Error"].to_numpy() / 100
                            vals = vals.tolist()
                            df_tab[idx_col] = vals
                    else:
                        t = (mat, lib_names_dict[idx_col[0]])
                        if idx_col[1] == "Value":
                            if mat != "H3":
                                vals = self.raw_data[t][4]["Value"].values[
                                    : len(x)
                                ]
                            else:
                                # Total activity
                                vals = []
                                for i in range(4):
                                    vals.extend(
                                        (
                                            self.raw_data[t][84]["Value"].values[
                                                i::4
                                            ]
                                        )
                                    )

                            df_tab[idx_col] = vals

                        elif idx_col[1] == "C/E Error":
                            if mat != "H3":
                                errs = self.raw_data[t][4]["Error"].values[
                                    : len(x)
                                ]
                            else:
                                errs = []
                                for i in range(4):
                                    yerr = self.raw_data[t][84]["Error"].values[
                                        i::4
                                    ]
                                    errs.extend(yerr)

                            vals1 = np.square(errs)
                            vals2 = np.square(
                                exp_data_df.loc[:, "Error"].to_numpy() / 100
                            )
                            ce_err = np.sqrt(vals1 + vals2)
                            ce_err = ce_err.tolist()
                            df_tab[idx_col] = ce_err
                        # Calculate C/E value
                        else:
                            if mat != "H3":
                                vals1 = self.raw_data[t][4]["Value"].values[
                                    : len(x)
                                ]
                            else:
                                vals1 = []
                                for i in range(4):
                                    vals1.extend(
                                        self.raw_data[t][84]["Value"].values[i::4]
                                    )

                            if mat == "H3":
                                vals2 = exp_data_df.loc[:, "Activity"].to_numpy()
                            else:
                                vals2 = exp_data_df.loc[:, "Reaction Rate"].to_numpy()
                            ratio = vals1 / vals2
                            ratio = ratio.tolist()
                            df_tab[idx_col] = vals1 / vals2

                # Assign worksheet title and put into Excel
                conv_df = self._get_conv_df(mat, len(x))
                sheet = self.testname.replace("-", " ")
                if mat != "H3":
                    sheet_name = sheet + ", Foil {}".format(mat)
                else:
                    sheet_name = sheet + " H3 activity"
                df_tab.to_excel(writer, sheet_name=sheet_name)
                conv_df.to_excel(writer, sheet_name=sheet_name, startrow=55)
                # Close the Pandas Excel writer object and output the Excel file

    def _build_atlas(self, tmp_path, atlas):
        """
        Build the Atlas (PDF) plots. See ExperimentalOutput documentation
        """
        code = "mcnp"
        for material in tqdm(self.inputs):
            # Tritium Activity
            if material == "H3":
                unit = "Bq/g"
                quantity = "Activity"
                for i in range(4):
                    data = []
                    # y = []
                    # err = []
                    exp_folder = os.path.join(self.path_exp_res, material)
                    exp_filename = self.testname + "_" + material + ".csv"
                    exp_filepath = os.path.join(exp_folder, exp_filename)
                    exp_data_df = pd.read_csv(exp_filepath)

                    xlabel = "Pellet no."
                    x = list(range(1, 13))

                    y = exp_data_df["Activity"].values[i * 12 : (i + 1) * 12]
                    err = exp_data_df["Error"].values[i * 12 : (i + 1) * 12] / 100

                    ylabel_exp = "Experiment"
                    data_exp = {"x": x, "y": y, "err": err, "ylabel": ylabel_exp}
                    data.append(data_exp)

                    for lib in self.lib[1:]:
                        # y = []
                        # err = []
                        # Total tritium production Li6 + Li7
                        ycalc = self.raw_data[(material, lib)][84][
                            "Value"
                        ].values[i::4]

                        yerr = np.square(
                            self.raw_data[(material, lib)][84]["Error"].values[
                                i::4
                            ]
                        )

                        y = ycalc
                        err = yerr

                        ylabel_calc = self.session.conf.get_lib_name(lib)
                        data_calc = {"x": x, "y": y, "err": err, "ylabel": ylabel_calc}
                        data.append(data_calc)

                    title = f"ENEA{2*(i+1)} pellet stack"
                    outname = "tmp"
                    plot = Plotter(
                        data,
                        title,
                        tmp_path,
                        outname,
                        quantity,
                        unit,
                        xlabel,
                        self.testname,
                    )
                    img_path = plot.plot("Discrete Experimental points")
                    atlas.insert_img(img_path)
            # Foils
            else:
                unit = "-"
                quantity = ["C/E"]
                data = []
                exp_folder = os.path.join(self.path_exp_res, material)
                exp_filename = self.testname + "_" + material + ".csv"
                exp_filepath = os.path.join(exp_folder, exp_filename)
                exp_data_df = pd.read_csv(exp_filepath)

                # Get experimental data and errors for the selected benchmark case
                xlabel = "Shielding thickness [cm]"
                x = list(exp_data_df["Depth"].values)
                y = []
                err = []
                y.append(exp_data_df["Reaction Rate"].values)
                err.append(exp_data_df["Error"].values / 100)
                # Append experimental data to data list (sent to plotter)
                ylabel = "Experiment"
                data_exp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                data.append(data_exp)

                title = self.testname + " experiment, Foil: " + material

                # Loop over selected libraries
                for lib in self.lib[1:]:
                    # Get library name, assign title to the plot
                    ylabel = self.session.conf.get_lib_name(lib)
                    y = []
                    err = []

                    ycalc = self.raw_data[(material, lib)][4]["Value"].values[
                        : len(x)
                    ]
                    y.append(ycalc)

                    yerr = self.raw_data[(material, lib)][4]["Error"].values[
                        : len(x)
                    ]
                    err.append(yerr)

                    # Append computational data to data list(to be sent to plotter)
                    data_comp = {"x": x, "y": y, "err": err, "ylabel": ylabel}
                    data.append(data_comp)

                outname = "tmp"
                plot = Plotter(
                    data,
                    title,
                    tmp_path,
                    outname,
                    quantity,
                    unit,
                    xlabel,
                    self.testname,
                )
                img_path = plot.plot("Waves")
                atlas.insert_img(img_path)
        return atlas

    def _get_conv_df(self, mat, size):
        conv_df = pd.DataFrame()
        code = "mcnp"
        for lib in self.lib[1:]:
            if mat != "H3":
                max = self.raw_data[(mat, lib)][4]["Error"].values[:size].max()
                avg = self.raw_data[(mat, lib)][4]["Error"].values[:size].mean()
            else:
                max = self.raw_data[(mat, lib)][84]["Error"].values[:size].max()
                avg = self.raw_data[(mat, lib)][84]["Error"].values[:size].mean()
            library = self.session.conf.get_lib_name(lib)
            conv_df.loc["Max Error", library] = max
            conv_df.loc["Average Error", library] = avg
        return conv_df
