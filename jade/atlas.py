# -*- coding: utf-8 -*-
"""
Created on Wed Jan 15 14:30:18 2020

@author: Davide Laghi

Copyright 2021, the JADE Development Team. All rights reserved.

This file is part of JADE.

JADE is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

JADE is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with JADE.  If not, see <http://www.gnu.org/licenses/>.
"""
import docx
# import win32com.client
import aspose.words
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
# from docx.shared import Pt
import os
import pandas as pd


class Atlas():
    def __init__(self, template, name):
        """
        Atlas of plots for post-processing

        Parameters
        ----------
        template : Path/str
            Word template for atlas
        # lib : list/str
        #     libraries to post-process

        Returns
        -------
        None.

        """
        # if type(lib) == list and len(lib) > 1:
        #     self.comparison = True
        #     name = ''
        #     for l in lib:
        #         name = name+l+'-'
        #     del[name[-1]]
        # else:
        #     self.comparison = False
        #     name = lib

        # self.lib = str(lib)  # Libraries to post process
        self.name = name  # Name of the Atlas (from libraries)
        # Open The Atlas template
        doc = docx.Document(template)
        doc.add_heading('JADE ATLAS: '+name, level=0)
        self.outname = 'atlas_' + name  # Name for the outfile
        self.doc = doc  # Word Document

    def insert_img(self, img, width=Inches(7.5)):
        self.doc.add_picture(img, width=width)
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def insert_df(self, df, caption=None, highlight=False, #, template_idx=None,
                  tablestyle=None):
        """
        Inser a dataframe as a table in a Word file

        Parameters
        ----------
        df : pd.DataFrame
            dataframe to insert.
        caption : str, optional
            caption of the table. The default is None
        highlight : bool, optional
            Very specific for stress assessment. The default is False.
        # template_idx : int, optional
        #     index of the template table to use. The default is None
        tablestyle : str, optional
            table word style to apply. The default is None

        Returns
        -------
        table : docx.Table
            table inserted.

        """
        # Create the table or inherit a template
        template_idx = None  # other possibilities not implemented anymore
        if template_idx is None:
            table = self.doc.add_table(1, len(df.columns))
        else:
            template = self.template_tables[template_idx]
            table = template.insert(self.doc)

        # Assign style if provided
        if table.style is not None:
            table.style = tablestyle

        # If template is not provided, the header row must be filled
        if template_idx is None:
            # add the header rows.
            for j in range(df.shape[-1]):
                table.cell(0, j).text = df.columns[j]

        # Add the rest of the data frame
        # val = df.values
        for i, (idx, row) in enumerate(df.iterrows()):
            # Understand is safety margin is barely acceptable
            flag_almost = False
            try:
                sm = float(row['Safety Margin'])
                if sm > 1 and sm < 1.1:
                    flag_almost = True
            except KeyError:
                pass
            except ValueError:
                pass
            except TypeError:
                # cannot convert to float
                pass

            row_cells = table.add_row().cells
            for j, item in enumerate(row):
                cell = row_cells[j]
                cell.text = str(item)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for par in cell.paragraphs:
                    par.style = 'Table'
                if highlight is not None:
                    if cell.text == 'NOK':
                        self._highlightCell(cell)
                    elif cell.text == 'OK' and flag_almost:
                        self._highlightCell(cell, color='FFFF46')

        if caption is not None:
            paragraph = self.doc.add_paragraph('Table ', style='Didascalia')
            self._wrapper(paragraph, 'table')
            paragraph.add_run(' - '+caption)
        # paragraph = doc.add_paragraph('Figure Text', style='Didascalia')

        return table

    def build(self, images_path, libmanager, mat_settings):
        """
        TO BE USED FOR SPHERE LEAKAGE BENCHMARK

        Generate the atlas filling the world template with plots and headings

        Parameters
        ----------
        images_path : str/path
            Path to temporary folder containig the images

        lib_manager : libmanager.LibManager
            Library manager for conversions and name recovery.

        mat_settings : pd.DataFrame
            contains settings for Materials

        Returns
        -------
        None.

        """
        # Build Atlas
        images = []
        tallies = []
        for img in os.listdir(images_path):
            img_path = os.path.join(images_path, img)
            pieces = img.split('-')
            zaid = pieces[0]
            tally = pieces[-1].split('.')[0]
            zaidnum = zaid.split('.')[0]
            if tally not in tallies:
                tallies.append(tally)

            images.append({'tally': tally, 'zaid': zaid, 'img': img_path,
                           'num': zaidnum})

        images = pd.DataFrame(images)
        # Reorder atlas
        images['num'] = pd.to_numeric(images['num'].values, errors='coerce')
        images.set_index('tally', inplace=True)
        for tally in tallies:
            self.doc.add_heading('Tally N.'+str(tally), level=1)
            # Be sure of the reordering
            df = images.loc[tally].sort_values('num')
            for idx, row in df.iterrows():
                title = 'Zaid: '+row['zaid']
                try:
                    name, formula = libmanager.get_zaidname(row['zaid'])
                    title = title+' ('+name+' '+formula+')'
                except ValueError:  # A material is passed instead of zaid
                    matname = mat_settings.loc[row['zaid'], 'Name']
                    title = title+' ('+matname+')'

                self.doc.add_heading(title, level=2)
                self.insert_img(row['img'])

    def save(self, outpath, pdfprint=True):
        """
        Save word atlas and possibly export PDF

        Parameters
        ----------
        outpath : str/path
            path to export the atlas
        pdfprint : Boolean, optional
            If True export also in PDF format

        Returns
        -------
        None.

        """
        outpath_word = os.path.join(outpath, self.outname+'.docx')
        outpath_pdf = os.path.join(outpath, self.outname+'.pdf')

        try:
            self.doc.save(outpath_word)
        except FileNotFoundError as e:
            print(' The following is the original exception:')
            print(e)
            print('\n it may be due to invalid characters in the file name')

        if pdfprint:
            in_file = outpath_word
            out_file = outpath_pdf

            doc = aspose.words.Document(in_file)
            doc.save(out_file)
            
#            word = win32com.client.Dispatch('Word.Application')
#            doc = word.Documents.Open(in_file)
#            doc.ExportAsFixedFormat(OutputFileName=out_file,
#                                    # 17 = PDF output, 18=XPS output
#                                    ExportFormat=17,
#                                    OpenAfterExport=False,
#                                    # 0=Print (higher res), 1=Screen (lower res)
#                                    OptimizeFor=0,
#   # 0=No bookmarks, 1=Heading bookmarks only, 2=bookmarks match word bookmarks
#                                    CreateBookmarks=1,
#                                    DocStructureTags=True)
#
#            doc.Close()
#            word.Quit()

    @staticmethod
    def _wrapper(paragraph, ptype):
        """
        Wrap a paragraph in order to add cross reference

        Parameters
        ----------
        paragraph : docx.Paragraph
            image to wrap.
        ptype : str
            type of paragraph to wrap

        Returns
        -------
        None.

        """
        if ptype == 'table':
            instruction = ' SEQ Table \\* ARABIC'
        elif ptype == 'figure':
            instruction = ' SEQ Figure \\* ARABIC'
        else:
            raise ValueError(ptype+' is not a supported paragraph type')

        run = run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = instruction
        r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)

    @staticmethod
    def _highlightCell(cell, color='FBD4B4'):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="'.format(nsdecls('w')) +
                                  color + r'"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm_1)
